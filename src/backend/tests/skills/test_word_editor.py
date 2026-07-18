"""End-to-end behavioural tests for the word-editing skill's editor engine.

Targets the rewritten editing surface:

  * P0-① per-run replace preserves formatting; cross-run match falls back
    to paragraph rebuild (the "edit destroys formatting" bug).
  * P0-② 0-replacement returns a ``warning`` so the LLM can self-correct.
  * P0-③ ``lenient=True`` matches across NBSP / zero-width / full-width.
  * P1-④ ``apply_edits`` runs an ordered op list in one open/save and
    reports per-op status (atomic semantics, gracefully handles invalid op).
  * P1-④ ``fill_placeholders`` + ``replace_many`` batch primitives.
  * P1-⑤ ``format_text`` selectors: anchor / style_filter / paragraph_indexes.
  * P2-⑦ ``replace_paragraph`` + ``delete_paragraph``.
  * P2-⑧ ``set_cell_text`` + ``fill_table``.
  * P2-⑨ ``list_placeholders``.
  * Run-merger: chained edits don't inflate the run count.

Runs in-process (no sandbox) — fast and deterministic. Uses the engine's
``_handle.use_workdir`` to pin the temp directory so concurrent tests don't
race on workdir state.
"""
from __future__ import annotations

import tempfile
from pathlib import Path

import pytest

from tests._skill_engine import load_engine

# Load the word-editing skill's vendored engine under a unique alias.
load_engine("word-editing", "word_engine")

from docx import Document  # noqa: E402

from word_engine._handle import use_workdir  # noqa: E402
from word_engine import editor  # noqa: E402

@pytest.fixture
def workdir(tmp_path: Path):
    with use_workdir(tmp_path):
        yield tmp_path

def _build_fixture_doc(path: Path) -> None:
    """Produce a doc with deliberately tricky structure:

    - Para 0: "foo" (bold) + "bar" (italic) — split runs mid-word.
    - Para 1: "hello world" — single bold run (per-run path target).
    - Para 2: "Dear {{name}}, your code is {{code}}." — placeholders.
    - Para 3: heading "背景".
    - Para 4: "草稿 — 待删除" — anchor for delete.
    - Para 5: "Total amount: 100 ​ yuan" — NBSP + ZWSP.
    - Para 6: "OLD PARAGRAPH WILL GO AWAY" — full-paragraph replace target.
    - One 2x3 table.
    """
    doc = Document()
    p0 = doc.add_paragraph()
    p0.add_run("foo").bold = True
    p0.add_run("bar").italic = True
    doc.add_paragraph().add_run("hello world").bold = True
    doc.add_paragraph("Dear {{name}}, your code is {{code}}.")
    doc.add_heading("背景", level=1)
    doc.add_paragraph("草稿 — 待删除")
    doc.add_paragraph("Total amount: 100 ​ yuan")
    doc.add_paragraph("OLD PARAGRAPH WILL GO AWAY")
    t = doc.add_table(rows=2, cols=3)
    t.cell(0, 0).text, t.cell(0, 1).text, t.cell(0, 2).text = "Name", "Q1", "Q2"
    t.cell(1, 0).text, t.cell(1, 1).text, t.cell(1, 2).text = "A", "100", "200"
    doc.save(str(path))

@pytest.fixture
def src_doc(workdir: Path) -> str:
    _build_fixture_doc(workdir / "in.docx")
    return "in.docx"

@pytest.fixture
def doc_with_table_in_middle(workdir: Path) -> str:
    """8 body paragraphs with a 3×2 table inserted in the middle (after the 5th paragraph).

    Body order: BodyA 0..4 (indices 0-4), [表格 R{r}C{c}], BodyB 0..2 (indices 5-7).
    → len(doc.paragraphs)==8, table_count==1, paragraph before the table at index 4 ("BodyA 4"),
    last paragraph at index 7 ("BodyB 2").
    (word_engine.reader/editor uses pure python-docx, no .NET dependency.)
    """
    doc = Document()
    for i in range(5):
        doc.add_paragraph(f"BodyA {i}")
    t = doc.add_table(rows=3, cols=2)
    for r in range(3):
        for c in range(2):
            t.cell(r, c).text = f"R{r}C{c}"
    for i in range(3):
        doc.add_paragraph(f"BodyB {i}")
    doc.save(workdir / "table_mid.docx")
    return "table_mid.docx"

# ── P0-① run-collapse fix ──────────────────────────────────────────────────

def test_per_run_replace_preserves_bold(workdir, src_doc):
    out = editor.search_and_replace(
        input_filename=src_doc, output_filename="t.docx",
        find="world", replace="WORLD",
    )
    assert out["replacements"] == 1
    d = Document(workdir / "t.docx")
    p1 = d.paragraphs[1]
    assert "WORLD" in p1.text
    # The hit was inside the single bold run; bold must survive.
    assert all(r.bold for r in p1.runs if r.text)

def test_replace_does_not_touch_unrelated_paragraph(workdir, src_doc):
    editor.search_and_replace(
        input_filename=src_doc, output_filename="t.docx",
        find="world", replace="WORLD",
    )
    d = Document(workdir / "t.docx")
    p0 = d.paragraphs[0]
    # Para 0 was [bold "foo", italic "bar"] and should remain split + styled.
    styled = [(r.text, bool(r.bold), bool(r.italic)) for r in p0.runs if r.text]
    assert ("foo", True, False) in styled
    assert ("bar", False, True) in styled

def test_cross_run_match_falls_back_to_paragraph_rebuild(workdir, src_doc):
    # 'foobar' lives across runs in para 0 → must still match (fallback path).
    out = editor.search_and_replace(
        input_filename=src_doc, output_filename="t.docx",
        find="foobar", replace="FOOBAR",
    )
    assert out["replacements"] == 1
    d = Document(workdir / "t.docx")
    assert "FOOBAR" in d.paragraphs[0].text

# ── P0-② warning on miss ───────────────────────────────────────────────────

def test_zero_replacements_returns_warning(workdir, src_doc):
    out = editor.search_and_replace(
        input_filename=src_doc, output_filename="t.docx",
        find="this string is not in the doc", replace="x",
    )
    assert out["replacements"] == 0
    assert "warning" in out
    assert "no matches" in out["warning"]

# ── P0-③ lenient match across NBSP + ZWSP ──────────────────────────────────

def test_strict_misses_nbsp(workdir, src_doc):
    out = editor.search_and_replace(
        input_filename=src_doc, output_filename="t.docx",
        find="Total amount: 100 yuan", replace="-",
    )
    assert out["replacements"] == 0

def test_lenient_matches_nbsp_and_zwsp(workdir, src_doc):
    out = editor.search_and_replace(
        input_filename=src_doc, output_filename="t.docx",
        find="Total amount: 100 yuan", replace="REPLACED", lenient=True,
    )
    assert out["replacements"] == 1
    d = Document(workdir / "t.docx")
    assert "REPLACED" in d.paragraphs[5].text

# ── P1-④ batch primitives ──────────────────────────────────────────────────

def test_fill_placeholders(workdir, src_doc):
    out = editor.fill_placeholders(
        input_filename=src_doc, output_filename="t.docx",
        mapping={"name": "Alice", "code": "9876"},
    )
    assert out["filled"] == {"name": 1, "code": 1}
    d = Document(workdir / "t.docx")
    assert d.paragraphs[2].text == "Dear Alice, your code is 9876."

def test_replace_many_with_partial_misses(workdir, src_doc):
    out = editor.replace_many(
        input_filename=src_doc, output_filename="t.docx",
        replacements=[
            {"find": "foo", "replace": "FOO"},
            {"find": "bar", "replace": "BAR"},
            {"find": "Q2", "replace": "Q-TWO"},
            {"find": "totally absent", "replace": "x"},
        ],
    )
    assert out["replacements"] == 3
    assert len(out["per_op"]) == 4
    assert out["per_op"][3]["replacements"] == 0
    assert "warning" in out["per_op"][3]

# ── P1-⑤ format_text selectors ─────────────────────────────────────────────

def test_format_text_by_anchor(workdir, src_doc):
    out = editor.format_text(
        input_filename=src_doc, output_filename="t.docx",
        anchor="背景", bold=True, color_hex="C00000",
    )
    assert out["paragraphs_touched"] == 1

def test_format_text_by_style_filter(workdir, src_doc):
    out = editor.format_text(
        input_filename=src_doc, output_filename="t.docx",
        style_filter="Heading", italic=True,
    )
    assert out["paragraphs_touched"] >= 1

def test_format_text_requires_one_selector(workdir, src_doc):
    with pytest.raises(ValueError, match="exactly one"):
        editor.format_text(
            input_filename=src_doc, output_filename="t.docx",
            paragraph_index=0, anchor="anything", bold=True,
        )

# ── format paragraph-level layout (line spacing, first-line indent, spacing) ─

def test_format_text_sets_line_spacing(workdir, src_doc):
    """``line_spacing=1.5`` writes ``paragraph_format.line_spacing = 1.5``
    (multiplier mode) onto every targeted paragraph.
    """
    out = editor.format_text(
        input_filename=src_doc, output_filename="t.docx",
        style_filter="Normal", line_spacing=1.5,
    )
    assert out["paragraphs_touched"] >= 1
    d = Document(workdir / "t.docx")
    sample = next(p for p in d.paragraphs if p.style and p.style.name == "Normal")
    assert sample.paragraph_format.line_spacing == 1.5

def test_format_text_first_line_indent_chars_uses_paragraph_font_size(workdir):
    """``first_line_indent_chars`` converts to Pt using the paragraph's
    effective font size. CJK convention: 1 char ≈ 1× font size.
    """
    from docx.shared import Pt
    src = "in.docx"
    d0 = Document()
    p = d0.add_paragraph("正文段落")
    # Set explicit run size so _para_effective_font_size_pt has a deterministic value.
    p.runs[0].font.size = Pt(12)
    d0.save(workdir / src)

    editor.format_text(
        input_filename=src, output_filename="t.docx",
        anchor="正文段落", first_line_indent_chars=2,
    )
    d = Document(workdir / "t.docx")
    target = next(p for p in d.paragraphs if "正文段落" in p.text)
    indent = target.paragraph_format.first_line_indent
    # 12pt × 2 chars = 24 Pt
    assert indent is not None
    assert abs(indent.pt - 24.0) < 0.01

def test_format_text_first_line_indent_chars_uses_concurrent_font_size(workdir):
    """When font_size is set in the SAME call, the indent uses THAT size, not
    the existing paragraph font size — so changing both at once stays
    visually consistent (e.g. switching to 小四/12pt + a 2-char first-line indent in one go).
    """
    from docx.shared import Pt
    src = "in.docx"
    d0 = Document()
    p = d0.add_paragraph("正文段落")
    p.runs[0].font.size = Pt(10)  # OLD font size
    d0.save(workdir / src)

    editor.format_text(
        input_filename=src, output_filename="t.docx",
        anchor="正文段落", font_size=12, first_line_indent_chars=2,
    )
    d = Document(workdir / "t.docx")
    target = next(p for p in d.paragraphs if "正文段落" in p.text)
    indent = target.paragraph_format.first_line_indent
    # 12pt × 2 = 24 Pt (NOT 10pt × 2 = 20)
    assert indent is not None
    assert abs(indent.pt - 24.0) < 0.01

def test_format_text_first_line_indent_pt_overrides_chars(workdir, src_doc):
    """When both ``first_line_indent_pt`` and ``_chars`` are passed, _pt wins."""
    out = editor.format_text(
        input_filename=src_doc, output_filename="t.docx",
        style_filter="Normal",
        first_line_indent_chars=2,
        first_line_indent_pt=30.0,
    )
    assert out["paragraphs_touched"] >= 1
    d = Document(workdir / "t.docx")
    sample = next(p for p in d.paragraphs if p.style and p.style.name == "Normal")
    assert sample.paragraph_format.first_line_indent is not None
    assert abs(sample.paragraph_format.first_line_indent.pt - 30.0) < 0.01

def test_format_text_space_before_after(workdir, src_doc):
    out = editor.format_text(
        input_filename=src_doc, output_filename="t.docx",
        style_filter="Normal",
        space_before_pt=6.0, space_after_pt=12.0,
    )
    assert out["paragraphs_touched"] >= 1
    d = Document(workdir / "t.docx")
    sample = next(p for p in d.paragraphs if p.style and p.style.name == "Normal")
    assert abs(sample.paragraph_format.space_before.pt - 6.0) < 0.01
    assert abs(sample.paragraph_format.space_after.pt - 12.0) < 0.01

def test_format_text_layout_only_no_run_field_required(workdir, src_doc):
    """Layout-only call (no bold/italic/font…) MUST be accepted — earlier
    versions required a run-level field which would block official-document (公文) formatting
    workflows that only need line_spacing + first_line_indent."""
    out = editor.format_text(
        input_filename=src_doc, output_filename="t.docx",
        style_filter="Normal",
        line_spacing=1.5, first_line_indent_chars=2,
    )
    assert out["paragraphs_touched"] >= 1

def test_format_text_font_name_sets_eastasia_for_cjk(workdir):
    """Regression for chat_20260509_164428: font_name must change the CJK
    eastAsia slot, not just the latin ascii slot. Otherwise Chinese
    characters keep rendering in the doc's previous CJK font and the
    user-visible font is unchanged.
    """
    from docx.shared import Pt
    from docx.oxml.ns import qn
    src = "in.docx"
    d0 = Document()
    p = d0.add_paragraph("中文正文段")
    p.runs[0].font.size = Pt(12)
    d0.save(workdir / src)

    editor.format_text(
        input_filename=src, output_filename="t.docx",
        anchor="中文正文段", font_name="仿宋",
    )
    d = Document(workdir / "t.docx")
    target = next(p for p in d.paragraphs if "中文正文段" in p.text)
    rPr = target.runs[0]._element.find(qn("w:rPr"))
    rFonts = rPr.find(qn("w:rFonts"))
    assert rFonts is not None
    # All three slots should carry the new font name now.
    assert rFonts.get(qn("w:ascii")) == "仿宋"
    assert rFonts.get(qn("w:hAnsi")) == "仿宋"
    assert rFonts.get(qn("w:eastAsia")) == "仿宋"

def test_format_text_style_filter_list_matches_any(workdir):
    """style_filter accepts a list of style names; matches any."""
    src = "in.docx"
    d0 = Document()
    p1 = d0.add_paragraph("Body Text para")
    p1.style = d0.styles["Body Text"]
    p2 = d0.add_paragraph("Normal para")  # default = Normal
    p3 = d0.add_paragraph()
    p3.add_run("Title para").bold = False
    p3.style = d0.styles["Title"]
    d0.save(workdir / src)

    out = editor.format_text(
        input_filename=src, output_filename="t.docx",
        style_filter=["Normal", "Body Text"], bold=True,
    )
    # Both body paragraphs touched, Title not touched.
    assert out["paragraphs_touched"] == 2
    d = Document(workdir / "t.docx")
    body_text_para = next(p for p in d.paragraphs if p.text == "Body Text para")
    normal_para = next(p for p in d.paragraphs if p.text == "Normal para")
    title_para = next(p for p in d.paragraphs if p.text == "Title para")
    assert body_text_para.runs[0].bold is True
    assert normal_para.runs[0].bold is True
    assert title_para.runs[0].bold is not True  # untouched

def test_format_text_style_filter_not_heading_matches_all_body(workdir):
    """``style_filter="!Heading"`` matches all non-heading paragraphs
    — covers Normal / Body Text / FirstParagraph / any other body style at
    once. Headings (Title / Heading N) are excluded.
    """
    src = "in.docx"
    d0 = Document()
    d0.add_heading("Chapter heading", level=1)  # Heading 1 → excluded
    d0.add_heading("Section heading", level=2)  # Heading 2 → excluded
    p = d0.add_paragraph("Body Text body")
    p.style = d0.styles["Body Text"]
    d0.add_paragraph("Normal body 1")
    d0.add_paragraph("Normal body 2")
    d0.save(workdir / src)

    out = editor.format_text(
        input_filename=src, output_filename="t.docx",
        style_filter="!Heading", italic=True,
    )
    assert out["paragraphs_touched"] == 3, out["indexes_touched"]
    d = Document(workdir / "t.docx")
    for p in d.paragraphs:
        if p.text in ("Body Text body", "Normal body 1", "Normal body 2"):
            assert p.runs[0].italic is True, p.text
        elif "heading" in p.text.lower():
            assert p.runs[0].italic is not True, p.text

def test_apply_edits_full_govdoc_layout(workdir):
    """End-to-end official-document layout: FangSong (仿宋) 12pt body + 1.5 line spacing +
    2-char first-line indent + HeiTi (黑体) 16pt headings + bold H2. Verifies the user-stated
    request works as a single apply_edits batch with the new fields.
    """
    from docx.shared import Pt
    src = "in.docx"
    d0 = Document()
    d0.add_heading("第一章 概述", level=1)
    d0.add_heading("1.1 背景", level=2)
    d0.add_paragraph("正文段一")
    d0.add_paragraph("正文段二")
    d0.add_heading("1.2 目标", level=2)
    d0.add_paragraph("正文段三")
    d0.save(workdir / src)

    out = editor.apply_edits(
        input_filename=src, output_filename="t.docx",
        ops=[
            {"op": "format", "style_filter": "Normal",
             "font_name": "方正仿宋简体", "font_size": 12,
             "line_spacing": 1.5, "first_line_indent_chars": 2},
            {"op": "format", "style_filter": "Heading 1",
             "font_name": "黑体", "font_size": 16},
            {"op": "format", "style_filter": "Heading 2", "bold": True},
        ],
    )
    assert out["ops_succeeded"] == 3, out

    d = Document(workdir / "t.docx")
    body_p = next(p for p in d.paragraphs if p.text == "正文段一")
    assert body_p.paragraph_format.line_spacing == 1.5
    # 12pt × 2 chars = 24 pt indent
    assert abs(body_p.paragraph_format.first_line_indent.pt - 24.0) < 0.01
    # Run-level: font name + size were applied
    assert body_p.runs[0].font.name == "方正仿宋简体"
    assert body_p.runs[0].font.size.pt == 12

    h1 = next(p for p in d.paragraphs if p.text == "第一章 概述")
    assert h1.runs[0].font.name == "黑体"
    assert h1.runs[0].font.size.pt == 16

    h2 = next(p for p in d.paragraphs if p.text == "1.1 背景")
    assert h2.runs[0].bold is True

# ── P2-⑦ paragraph-level ops ───────────────────────────────────────────────

# NOTE: replace_paragraph / delete_paragraph / set_cell_text / fill_table are no
# longer exposed as standalone engine functions. They live on as ops inside
# ``apply_edits`` (``_OP_REGISTRY``); the tests below exercise them via that
# path so we keep coverage of the underlying primitives.

def test_replace_paragraph_via_apply_edits(workdir, src_doc):
    editor.apply_edits(
        input_filename=src_doc, output_filename="t.docx",
        ops=[{"op": "replace_paragraph", "anchor": "OLD PARAGRAPH",
              "new_text": "THIS IS THE NEW PARA"}],
    )
    d = Document(workdir / "t.docx")
    assert any(p.text == "THIS IS THE NEW PARA" for p in d.paragraphs)

def test_replace_paragraph_auto_renders_markdown(workdir, src_doc):
    """replace_paragraph with markdown content renders real Word blocks (auto)."""
    out = editor.apply_edits(
        input_filename=src_doc, output_filename="t.docx",
        ops=[{"op": "replace_paragraph", "anchor": "OLD PARAGRAPH",
              "new_text": "### 新标题\n\n**重点**说明\n\n- 第一条\n- 第二条"}],
    )
    res = out["results"][0]
    assert res["ok"] and res["format"] == "markdown"
    d = Document(workdir / "t.docx")
    # literal '### ' markup must NOT survive; heading rendered as a real style
    assert not any("### " in p.text for p in d.paragraphs)
    assert any(
        "新标题" in p.text and (p.style.name or "").startswith("Heading")
        for p in d.paragraphs
    )
    # bullet list items present and old paragraph gone
    assert any(p.text == "第一条" for p in d.paragraphs)
    assert not any("OLD PARAGRAPH" in p.text for p in d.paragraphs)

def test_replace_section_auto_renders_markdown(workdir, src_doc):
    """replace_section dumps a markdown draft → headings/lists rendered, old
    body removed, original heading preserved."""
    out = editor.apply_edits(
        input_filename=src_doc, output_filename="t.docx",
        ops=[{"op": "replace_section", "heading_anchor": "背景",
              "new_content": "#### 小节一\n\n正文**强调**内容\n\n- 要点 A\n- 要点 B"}],
    )
    res = out["results"][0]
    assert res["ok"] and res["format"] == "markdown"
    d = Document(workdir / "t.docx")
    texts = [p.text for p in d.paragraphs]
    assert "背景" in texts                       # heading preserved
    assert not any("草稿" in t for t in texts)    # old section body removed
    assert not any("#### " in t for t in texts)   # markdown rendered, not literal
    assert any("小节一" in t for t in texts)
    assert any(t == "要点 A" for t in texts)

def test_replace_section_inherits_indent_and_drops_blank_lines(workdir):
    """Repro: a doc whose body paras carry firstLineChars=200 (official-document 2-char indent). After
    replace_section with \\n\\n-separated prose, the new paras must (a) keep the
    same first-line indent and (b) NOT leave empty paragraphs between them."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    d = Document()
    d.add_heading("一、概述", level=1)
    p = d.add_paragraph("原有正文一，带两格首行缩进。")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:firstLineChars"), "200")
    p._element.get_or_add_pPr().append(ind)
    d.add_paragraph("原有正文二。")
    d.add_heading("二、下一节", level=1)
    d.save(str(workdir / "in.docx"))

    out = editor.apply_edits(
        input_filename="in.docx", output_filename="t.docx",
        ops=[{"op": "replace_section", "heading_anchor": "一、概述",
              "new_content": "新正文第一段。\n\n新正文第二段。\n\n新正文第三段。"}],
    )
    res = out["results"][0]
    assert res["ok"], res
    assert res["format"] == "text"
    assert res["new_paragraph_count"] == 3  # \n\n separators, NOT 5 with blanks

    rd = Document(workdir / "t.docx")
    new_bodies = [p for p in rd.paragraphs if p.text.startswith("新正文")]
    assert len(new_bodies) == 3
    # no empty paragraph slipped in between the three
    assert all(p.text.strip() for p in new_bodies)
    # every new paragraph kept the 2-char first-line indent
    for p in new_bodies:
        pPr = p._element.find(qn("w:pPr"))
        ind = pPr.find(qn("w:ind")) if pPr is not None else None
        assert ind is not None and ind.get(qn("w:firstLineChars")) == "200", (
            f"missing inherited indent on {p.text!r}"
        )

def test_replace_section_markdown_table_failure_is_atomic(workdir, src_doc):
    """A markdown table inside new_content raises (use add_table) but the old
    section content must be left intact — emit happens before deletion."""
    out = editor.apply_edits(
        input_filename=src_doc, output_filename="t.docx",
        ops=[{"op": "replace_section", "heading_anchor": "背景",
              "new_content": "正文\n\n| a | b |\n|---|---|\n| 1 | 2 |"}],
    )
    res = out["results"][0]
    assert not res["ok"]
    assert "table" in res["error"].lower()
    # section untouched: original body still present, no orphan paragraphs leaked
    d = Document(workdir / "t.docx")
    texts = [p.text for p in d.paragraphs]
    assert any("草稿" in t for t in texts)
    assert not any(t.strip() == "正文" for t in texts)

def test_insert_image_via_image_path(workdir, src_doc):
    """insert_image accepts a direct sandbox path (image_path) — no --image alias."""
    import struct
    import zlib

    def _chunk(typ, data):
        body = typ + data
        return struct.pack(">I", len(data)) + body + struct.pack(
            ">I", zlib.crc32(body) & 0xFFFFFFFF
        )

    png = (
        b"\x89PNG\r\n\x1a\n"
        + _chunk(b"IHDR", struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0))
        + _chunk(b"IDAT", zlib.compress(b"\x00\xff\x00\x00"))
        + _chunk(b"IEND", b"")
    )
    img = workdir / "chart1.png"
    img.write_bytes(png)
    out = editor.apply_edits(
        input_filename=src_doc, output_filename="t.docx",
        ops=[{"op": "insert_image", "image_path": str(img),
              "position": "end", "width_cm": 5}],
    )
    res = out["results"][0]
    assert res["ok"], res
    assert res["image_source"] == str(img)
    from docx.oxml.ns import qn
    blips = Document(workdir / "t.docx").element.body.findall(".//" + qn("a:blip"))
    assert len(blips) >= 1  # an inline picture got embedded

def test_insert_image_requires_a_source(workdir, src_doc):
    """Neither image_path nor image_filename → clear error, batch survives."""
    out = editor.apply_edits(
        input_filename=src_doc, output_filename="t.docx",
        ops=[{"op": "insert_image", "position": "end"}],
    )
    res = out["results"][0]
    assert not res["ok"]
    assert "image_path" in res["error"]

def test_insert_image_missing_file_reports_clearly(workdir, src_doc):
    """image_path pointing nowhere → FileNotFound with a put-into-sandbox hint."""
    out = editor.apply_edits(
        input_filename=src_doc, output_filename="t.docx",
        ops=[{"op": "insert_image", "image_path": "/workspace/nope.png",
              "position": "end"}],
    )
    res = out["results"][0]
    assert not res["ok"]
    assert "not found" in res["error"].lower()

def test_delete_paragraph_via_apply_edits(workdir, src_doc):
    n_before = len(Document(workdir / src_doc).paragraphs)
    editor.apply_edits(
        input_filename=src_doc, output_filename="t.docx",
        ops=[{"op": "delete_paragraph", "anchor": "草稿"}],
    )
    n_after = len(Document(workdir / "t.docx").paragraphs)
    assert n_after == n_before - 1

# ── P2-⑧ table ops (via apply_edits) ───────────────────────────────────────

def test_set_cell_text_via_apply_edits(workdir, src_doc):
    editor.apply_edits(
        input_filename=src_doc, output_filename="t.docx",
        ops=[{"op": "set_cell_text", "table_index": 0, "row": 1, "col": 1, "text": "999"}],
    )
    d = Document(workdir / "t.docx")
    assert d.tables[0].cell(1, 1).text == "999"

def test_fill_table_append_via_apply_edits(workdir, src_doc):
    editor.apply_edits(
        input_filename=src_doc, output_filename="t.docx",
        ops=[{"op": "fill_table", "table_index": 0,
              "rows": [["B", "300", "400"], ["C", "500", "600"]], "mode": "append"}],
    )
    d = Document(workdir / "t.docx")
    assert len(d.tables[0].rows) == 4

def test_fill_table_overwrite_keeps_header_via_apply_edits(workdir, src_doc):
    editor.apply_edits(
        input_filename=src_doc, output_filename="t.docx",
        ops=[{"op": "fill_table", "table_index": 0,
              "rows": [["X", "1", "2"]], "mode": "overwrite", "has_header": True}],
    )
    d = Document(workdir / "t.docx")
    assert len(d.tables[0].rows) == 2
    assert d.tables[0].cell(1, 0).text == "X"

# ── add_table at arbitrary positions ───────────────────────────────────────
#
# Verifies the mid-document table insertion paths via the ``add_table`` op
# in ``apply_edits``. Each test reads the resulting body XML in document
# order and asserts the new <w:tbl> ends up at the expected location.

def _body_block_order(doc) -> list[tuple[str, str]]:
    """Return [(kind, label), …] for each top-level body element in order.
    kind = 'p' (paragraph) | 'tbl' (table); label = first-cell text or para text.
    """
    from docx.oxml.ns import qn
    out: list[tuple[str, str]] = []
    body = doc.element.body
    for child in body:
        if child.tag == qn("w:p"):
            text = "".join(t.text or "" for t in child.iter(qn("w:t")))
            out.append(("p", text))
        elif child.tag == qn("w:tbl"):
            # First cell's text as a marker
            first_cell = child.find(f"{qn('w:tr')}/{qn('w:tc')}")
            label = ""
            if first_cell is not None:
                label = "".join(t.text or "" for t in first_cell.iter(qn("w:t")))
            out.append(("tbl", label))
    return out

def test_add_table_after_heading_via_apply_edits(workdir, src_doc):
    editor.apply_edits(
        input_filename=src_doc, output_filename="t.docx",
        ops=[{
            "op": "add_table",
            "rows": [["X", "Y"], ["1", "2"]],
            "has_header": True,
            "position": "after_heading",
            "anchor": "背景",
        }],
    )
    d = Document(workdir / "t.docx")
    blocks = _body_block_order(d)
    # Find the heading row, then the next non-empty block must be the new tbl
    h_idx = next(i for i, (kind, label) in enumerate(blocks) if kind == "p" and "背景" in label)
    next_kind, next_label = blocks[h_idx + 1]
    assert next_kind == "tbl"
    assert next_label == "X"
    # And the original first table (Name|Q1|Q2) must still exist somewhere later
    assert any(kind == "tbl" and label == "Name" for kind, label in blocks)

def test_add_table_before_paragraph_via_apply_edits(workdir, src_doc):
    editor.apply_edits(
        input_filename=src_doc, output_filename="t.docx",
        ops=[{
            "op": "add_table",
            "rows": [["TOP_TABLE_HEADER"], ["row1"]],
            "has_header": True,
            "position": "before_paragraph",
            "anchor": "OLD PARAGRAPH",
        }],
    )
    d = Document(workdir / "t.docx")
    blocks = _body_block_order(d)
    target_idx = next(i for i, (kind, label) in enumerate(blocks)
                      if kind == "p" and "OLD PARAGRAPH" in label)
    prev_kind, prev_label = blocks[target_idx - 1]
    assert prev_kind == "tbl"
    assert prev_label == "TOP_TABLE_HEADER"

def test_add_table_with_caption_after_paragraph_via_apply_edits(workdir, src_doc):
    editor.apply_edits(
        input_filename=src_doc, output_filename="t.docx",
        ops=[{
            "op": "add_table",
            "rows": [["CAP_TABLE"], ["v"]],
            "has_header": True,
            "caption": "表 X. 测试用表格",
            "position": "after_paragraph",
            "anchor": "草稿",
        }],
    )
    d = Document(workdir / "t.docx")
    blocks = _body_block_order(d)
    # Layout right after the anchor: anchor → caption → tbl
    a_idx = next(i for i, (kind, label) in enumerate(blocks)
                 if kind == "p" and "草稿" in label)
    assert blocks[a_idx + 1] == ("p", "表 X. 测试用表格")
    assert blocks[a_idx + 2][0] == "tbl"
    assert blocks[a_idx + 2][1] == "CAP_TABLE"

def test_add_table_at_start_via_apply_edits(workdir, src_doc):
    editor.apply_edits(
        input_filename=src_doc, output_filename="t.docx",
        ops=[{
            "op": "add_table",
            "rows": [["FIRST_TBL"], ["x"]],
            "has_header": True,
            "position": "start",
        }],
    )
    d = Document(workdir / "t.docx")
    blocks = _body_block_order(d)
    assert blocks[0] == ("tbl", "FIRST_TBL")

def test_add_table_anchor_required_for_after_heading(workdir, src_doc):
    out = editor.apply_edits(
        input_filename=src_doc, output_filename="t.docx",
        ops=[{
            "op": "add_table",
            "rows": [["a"], ["b"]],
            "position": "after_heading",  # missing anchor
        }],
    )
    assert out["ops_failed"] == 1
    assert "anchor" in out["results"][0]["error"]

def test_add_table_unknown_anchor_fails_gracefully(workdir, src_doc):
    out = editor.apply_edits(
        input_filename=src_doc, output_filename="t.docx",
        ops=[{
            "op": "add_table",
            "rows": [["a"], ["b"]],
            "position": "after_heading",
            "anchor": "UNKNOWN HEADING",
        }],
    )
    assert out["ops_failed"] == 1
    assert "no heading containing" in out["results"][0]["error"]

# ── P2-⑨ discovery ─────────────────────────────────────────────────────────

def test_list_placeholders(workdir, src_doc):
    out = editor.list_placeholders(input_filename=src_doc)
    assert set(out["placeholders"]) == {"name", "code"}
    assert out["counts"] == {"name": 1, "code": 1}

# ── P1-④ apply_edits — the headline tool ───────────────────────────────────

def test_apply_edits_atomic_with_partial_failure(workdir, src_doc):
    ops = [
        {"op": "fill_placeholders", "mapping": {"name": "Bob", "code": "1234"}},
        {"op": "format", "anchor": "背景", "bold": True, "color_hex": "C00000"},
        {"op": "delete_paragraph", "anchor": "草稿"},
        {"op": "replace_paragraph", "anchor": "OLD PARAGRAPH", "new_text": "REPLACED PARA"},
        {"op": "set_cell_text", "table_index": 0, "row": 1, "col": 0, "text": "Z"},
        {"op": "fill_table", "table_index": 0, "rows": [["Q", "9", "9"]], "mode": "append"},
        {"op": "insert", "position": "after_heading", "anchor": "背景", "text": "新增段落"},
        {"op": "replace", "find": "foo", "replace": "FOO"},
        {"op": "totally_invalid_op", "x": 1},
    ]
    out = editor.apply_edits(
        input_filename=src_doc, output_filename="t.docx", ops=ops,
    )
    assert out["ops_total"] == 9
    assert out["ops_succeeded"] == 8
    assert out["ops_failed"] == 1
    assert not out["results"][8]["ok"]
    assert "unknown op" in out["results"][8]["error"]
    d = Document(workdir / "t.docx")
    assert "Bob" in d.paragraphs[2].text and "1234" in d.paragraphs[2].text
    assert any("REPLACED PARA" in p.text for p in d.paragraphs)
    assert any("新增段落" in p.text for p in d.paragraphs)
    assert d.tables[0].cell(1, 0).text == "Z"
    assert len(d.tables[0].rows) == 3  # header + A + Q

def test_apply_edits_stop_on_error(workdir, src_doc):
    out = editor.apply_edits(
        input_filename=src_doc, output_filename="t.docx",
        stop_on_error=True,
        ops=[
            {"op": "replace", "find": "foo", "replace": "FOO"},
            {"op": "delete_paragraph", "anchor": "totally absent"},
            {"op": "replace", "find": "bar", "replace": "BAR"},  # should NOT run
        ],
    )
    assert out["ops_succeeded"] == 1
    assert out["ops_failed"] == 1
    # Third op never executed because we stopped.
    assert len(out["results"]) == 2

# ── run merger sanity ──────────────────────────────────────────────────────

def test_run_merger_keeps_run_count_bounded(workdir, src_doc):
    cur = src_doc
    for i in range(5):
        nxt = f"step_{i}.docx"
        editor.search_and_replace(
            input_filename=cur, output_filename=nxt,
            find="foo" if i % 2 == 0 else "FOO",
            replace="FOO" if i % 2 == 0 else "foo",
        )
        cur = nxt
    d = Document(workdir / cur)
    n_runs_p0 = len([r for r in d.paragraphs[0].runs if r.text])
    # Para 0 originally has 2 runs; without merging this would inflate to ≥6.
    assert n_runs_p0 <= 3, f"run count exploded: {n_runs_p0}"

# ── New: multi-paragraph cell content (P0a — fixes "PDF table garbled text") ─

def _cell_paragraph_count(cell) -> int:
    return len(list(cell.paragraphs))

def test_set_cell_text_multiline_creates_multi_paragraphs(workdir, src_doc):
    """Verify ``\\n`` inside cell value becomes multiple <w:p>, not literal LF."""
    editor.apply_edits(
        input_filename=src_doc, output_filename="t.docx",
        ops=[{"op": "set_cell_text", "table_index": 0, "row": 1, "col": 1,
              "text": "• line A\n• line B\n• line C"}],
    )
    d = Document(workdir / "t.docx")
    cell = d.tables[0].cell(1, 1)
    assert _cell_paragraph_count(cell) == 3
    para_texts = [p.text for p in cell.paragraphs]
    assert para_texts == ["• line A", "• line B", "• line C"]
    # Critical: no literal LF should remain in any <w:t>
    for p in cell.paragraphs:
        for r in p.runs:
            assert "\n" not in r.text, f"literal LF survived in run: {r.text!r}"

def test_fill_table_multiline_cell_via_apply_edits(workdir, src_doc):
    editor.apply_edits(
        input_filename=src_doc, output_filename="t.docx",
        ops=[{"op": "fill_table", "table_index": 0, "mode": "append",
              "rows": [["X", "single", "• L1\n• L2\n• L3"]]}],
    )
    d = Document(workdir / "t.docx")
    new_row = d.tables[0].rows[2]
    assert new_row.cells[0].text == "X"
    assert _cell_paragraph_count(new_row.cells[2]) == 3

def test_add_table_multiline_cells_via_apply_edits(workdir, src_doc):
    editor.apply_edits(
        input_filename=src_doc, output_filename="t.docx",
        ops=[{"op": "add_table",
              "position": "end",
              "rows": [
                  ["序号", "要点"],
                  ["一", "• point A\n• point B"],
                  ["二", "• X\n• Y\n• Z"],
              ],
              "has_header": True}],
    )
    d = Document(workdir / "t.docx")
    # The new table is the LAST one in the doc
    new_tbl = d.tables[-1]
    assert _cell_paragraph_count(new_tbl.cell(1, 1)) == 2
    assert _cell_paragraph_count(new_tbl.cell(2, 1)) == 3
    # No literal LF anywhere in the new table
    from docx.oxml.ns import qn
    all_w_t = list(new_tbl._element.iter(qn("w:t")))
    assert all("\n" not in (t.text or "") for t in all_w_t), "LF leaked into <w:t>"

# ── New: replace_paragraph with multi-line new_text (P0b) ───────────────────

def test_replace_paragraph_multiline_splits_into_multiple_paragraphs(workdir, src_doc):
    n_before = len(Document(workdir / src_doc).paragraphs)
    out = editor.apply_edits(
        input_filename=src_doc, output_filename="t.docx",
        ops=[{"op": "replace_paragraph", "anchor": "OLD PARAGRAPH",
              "new_text": "Para A\nPara B\nPara C"}],
    )
    assert out["ops_succeeded"] == 1
    res = out["results"][0]
    assert res["new_paragraph_count"] == 3
    d = Document(workdir / "t.docx")
    # 3 new paragraphs replace 1 old → +2 net
    assert len(d.paragraphs) == n_before + 2
    texts = [p.text for p in d.paragraphs]
    assert "Para A" in texts and "Para B" in texts and "Para C" in texts
    # The old paragraph text should be gone
    assert not any("OLD PARAGRAPH WILL GO AWAY" in t for t in texts)

# ── New: replace_section (P1a — atomic section rewrite) ─────────────────────

def _build_doc_with_section(path: Path) -> None:
    """Build a doc with a multi-paragraph H2 section to be replaced."""
    doc = Document()
    doc.add_heading("第一章 引言", level=1)
    doc.add_paragraph("引言段落 1")
    doc.add_heading("1.1 背景", level=2)
    doc.add_paragraph("背景段落 1")
    doc.add_paragraph("背景段落 2")
    doc.add_paragraph("背景段落 3")
    doc.add_heading("1.2 目标", level=2)
    doc.add_paragraph("目标段落")
    doc.add_heading("第二章 方法", level=1)
    doc.add_paragraph("方法段落")
    doc.save(str(path))

@pytest.fixture
def section_doc(workdir: Path) -> str:
    _build_doc_with_section(workdir / "section.docx")
    return "section.docx"

def test_replace_section_preserve_heading(workdir, section_doc):
    out = editor.apply_edits(
        input_filename=section_doc, output_filename="t.docx",
        ops=[{"op": "replace_section",
              "heading_anchor": "1.1 背景",
              "new_content": "新背景段落 A\n新背景段落 B"}],
    )
    assert out["ops_succeeded"] == 1
    d = Document(workdir / "t.docx")
    texts = [p.text for p in d.paragraphs]
    # Heading kept
    assert "1.1 背景" in texts
    # Original 3 body paras gone
    assert "背景段落 1" not in texts
    assert "背景段落 2" not in texts
    assert "背景段落 3" not in texts
    # New body inserted
    assert "新背景段落 A" in texts
    assert "新背景段落 B" in texts
    # Sibling 1.2 must NOT have been touched (proves boundary detection works)
    assert "1.2 目标" in texts
    assert "目标段落" in texts

def test_replace_section_stops_at_higher_level_heading(workdir, section_doc):
    """Replacing a Heading 1 section must NOT spill across into the next H1."""
    out = editor.apply_edits(
        input_filename=section_doc, output_filename="t.docx",
        ops=[{"op": "replace_section",
              "heading_anchor": "第一章 引言",
              "new_content": "全新引言"}],
    )
    assert out["ops_succeeded"] == 1
    d = Document(workdir / "t.docx")
    texts = [p.text for p in d.paragraphs]
    # Chapter 2 untouched
    assert "第二章 方法" in texts
    assert "方法段落" in texts
    # Original body of chapter 1 (incl. 1.1 / 1.2 sub-headings) all gone
    assert "1.1 背景" not in texts
    assert "1.2 目标" not in texts
    assert "新全新引言" not in texts and "全新引言" in texts

def test_replace_section_unknown_heading_fails_gracefully(workdir, section_doc):
    out = editor.apply_edits(
        input_filename=section_doc, output_filename="t.docx",
        ops=[{"op": "replace_section",
              "heading_anchor": "不存在的章节",
              "new_content": "x"}],
    )
    assert out["ops_failed"] == 1
    assert "no heading containing" in out["results"][0]["error"]

# ── New: delete_range op (P1a) ──────────────────────────────────────────────

def test_delete_range_atomic(workdir, section_doc):
    n_before = len(Document(workdir / section_doc).paragraphs)
    out = editor.apply_edits(
        input_filename=section_doc, output_filename="t.docx",
        ops=[{"op": "delete_range",
              "start_anchor": "背景段落 1",
              "end_anchor": "背景段落 3",
              "include_end": True}],
    )
    assert out["ops_succeeded"] == 1
    d = Document(workdir / "t.docx")
    texts = [p.text for p in d.paragraphs]
    assert "背景段落 1" not in texts
    assert "背景段落 2" not in texts
    assert "背景段落 3" not in texts
    # exactly 3 deleted
    assert len(d.paragraphs) == n_before - 3

# ── New: int-anchor delete_paragraph drift defence (P1b) ────────────────────

def test_apply_edits_int_anchor_deletes_no_drift(workdir):
    """Sequential int-anchored delete_paragraph ops must NOT skip every other
    paragraph. The runtime auto-reorders them descending to avoid drift.
    """
    # Build a doc with 6 cleanly numbered paragraphs.
    doc = Document()
    for i in range(6):
        doc.add_paragraph(f"P{i}")
    src = "drift.docx"
    doc.save(workdir / src)

    # Ask to delete paragraphs 1, 2, 3 by int index in input order.
    # If executed naively in order, indexes drift and only P1, P3, P5 get deleted.
    # With the reorder defence, P3 → P2 → P1 are deleted (descending), leaving P0, P4, P5.
    out = editor.apply_edits(
        input_filename=src, output_filename="t.docx",
        ops=[
            {"op": "delete_paragraph", "anchor": 1},
            {"op": "delete_paragraph", "anchor": 2},
            {"op": "delete_paragraph", "anchor": 3},
        ],
    )
    assert out["ops_succeeded"] == 3
    assert out.get("int_anchor_deletes_reordered") == 3
    d = Document(workdir / "t.docx")
    texts = [p.text for p in d.paragraphs]
    assert texts == ["P0", "P4", "P5"], texts

def test_apply_edits_text_anchor_deletes_left_alone(workdir, section_doc):
    """Text-anchor delete_paragraph ops are NOT reordered (drift-immune already)."""
    out = editor.apply_edits(
        input_filename=section_doc, output_filename="t.docx",
        ops=[
            {"op": "delete_paragraph", "anchor": "背景段落 1"},
            {"op": "delete_paragraph", "anchor": "背景段落 2"},
        ],
    )
    assert out["ops_succeeded"] == 2
    assert "int_anchor_deletes_reordered" not in out
    d = Document(workdir / "t.docx")
    texts = [p.text for p in d.paragraphs]
    assert "背景段落 1" not in texts
    assert "背景段落 2" not in texts

# ── New: cross-paragraph find warning (P2) ──────────────────────────────────

def test_replace_text_with_newline_in_find_returns_specific_warning(workdir, src_doc):
    out = editor.search_and_replace(
        input_filename=src_doc, output_filename="t.docx",
        find="hello world\nDear", replace="x",
    )
    assert out["replacements"] == 0
    assert "warning" in out
    # The warning must point to the cross-paragraph issue specifically.
    assert "newline" in out["warning"].lower() or "\\n" in out["warning"]
    assert "replace_section" in out["warning"]

# ── New: JSON-string arg coercion at MCP wrapper layer (P0c) ────────────────

def test_reader_get_text_paragraph_count_matches_body_only(workdir, doc_with_table_in_middle):
    """The fix: paragraph_count emitted by get_text MUST match
    ``len(doc.paragraphs)`` (body only), not body + table-cell paragraphs.
    """
    from word_engine import reader

    out = reader.get_text(input_filename=doc_with_table_in_middle)
    d = Document(workdir / doc_with_table_in_middle)
    assert out["paragraph_count"] == len(d.paragraphs) == 8
    assert out["table_count"] == 1
    # The text must contain ONLY body paragraphs, NOT cells like 'R0C0'.
    assert "R0C0" not in out["text"]
    assert "BodyA 0" in out["text"]
    assert "BodyB 2" in out["text"]

def test_reader_get_text_range_indexes_align_with_editor(workdir, doc_with_table_in_middle):
    """An index returned by get_text MUST be addressable by edit ops on the
    same doc. This is the regression: previously index 12 was valid for
    get_text (counts cells) but invalid for delete_paragraph (body-only).
    """
    from word_engine import reader

    out = reader.get_text(
        input_filename=doc_with_table_in_middle,
        paragraph_range=(7, 8),  # last body paragraph
    )
    assert out["selected_range"] == [7, 8]
    assert out["text"] == "BodyB 2"

    # Now delete the same paragraph via apply_edits using the same index.
    res = editor.apply_edits(
        input_filename=doc_with_table_in_middle, output_filename="t.docx",
        ops=[{"op": "delete_paragraph", "anchor": 7}],
    )
    assert res["ops_succeeded"] == 1
    d = Document(workdir / "t.docx")
    assert "BodyB 2" not in [p.text for p in d.paragraphs]

def test_reader_get_text_rejects_out_of_range_with_clear_count(
    workdir, doc_with_table_in_middle
):
    """Asking for a range past body length should fail cleanly (the new error
    message names the BODY paragraph count, not the inflated body+cells count).
    """
    from word_engine import reader

    with pytest.raises(ValueError, match="8 body paragraphs"):
        reader.get_text(
            input_filename=doc_with_table_in_middle,
            paragraph_range=(0, 99),
        )

def test_outline_includes_tables_with_indexes(workdir, doc_with_table_in_middle):
    """The new ``tables`` field on outline lets the LLM locate tables by index
    + spatial position (preceding paragraph) without scanning text.
    """
    from word_engine import reader

    out = reader.get_outline(input_filename=doc_with_table_in_middle)
    assert out["table_count"] == 1
    assert out["paragraph_count"] == 8
    tbl = out["tables"][0]
    assert tbl["table_index"] == 0
    assert tbl["rows"] == 3 and tbl["cols"] == 2
    assert tbl["first_cell_text"] == "R0C0"
    # Table sits after body paragraph 4 ("BodyA 4")
    assert tbl["preceding_paragraph_index"] == 4
    assert tbl["preceding_paragraph_text"] == "BodyA 4"

# ── move_table op ──────────────────────────────────────────────────────────

def test_move_table_to_after_heading(workdir):
    """Build doc: heading, body, body, body, heading, body, TABLE_AT_END.
    Move the table to right after the first heading.
    """
    doc = Document()
    doc.add_heading("第一章", level=1)
    doc.add_paragraph("引言段落")
    doc.add_heading("结论", level=1)
    doc.add_paragraph("结论段落")
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "MOVE_ME_HEADER"
    t.cell(0, 1).text = "x"
    t.cell(1, 0).text = "row1"
    t.cell(1, 1).text = "y"
    src = "movetbl.docx"
    doc.save(workdir / src)

    res = editor.apply_edits(
        input_filename=src, output_filename="out.docx",
        ops=[{"op": "move_table", "table_index": 0,
              "position": "after_heading", "anchor": "第一章"}],
    )
    assert res["ops_succeeded"] == 1
    res_op = res["results"][0]
    assert res_op["ok"] is True

    # Verify by walking body-level XML in document order.
    d = Document(workdir / "out.docx")
    from docx.oxml.ns import qn
    body = d.element.body
    seq = []
    for child in body:
        if child.tag == qn("w:p"):
            seq.append(("p", "".join(t.text or "" for t in child.iter(qn("w:t")))))
        elif child.tag == qn("w:tbl"):
            first_cell = child.find(f"{qn('w:tr')}/{qn('w:tc')}")
            label = "".join(t.text or "" for t in first_cell.iter(qn("w:t"))) if first_cell is not None else ""
            seq.append(("tbl", label))
    # Find the heading "第一章", the table must be the next thing.
    h_idx = next(i for i, (k, l) in enumerate(seq) if k == "p" and l == "第一章")
    assert seq[h_idx + 1][0] == "tbl"
    assert seq[h_idx + 1][1] == "MOVE_ME_HEADER"

def test_move_table_invalid_index_fails_gracefully(workdir, doc_with_table_in_middle):
    res = editor.apply_edits(
        input_filename=doc_with_table_in_middle, output_filename="t.docx",
        ops=[{"op": "move_table", "table_index": 99, "position": "start"}],
    )
    assert res["ops_failed"] == 1
    assert "out of range" in res["results"][0]["error"]

# ── auto vertical-merge of empty cells ─────────────────────────────────────

def _vmerge_at(table, row_idx: int, col_idx: int):
    """Return the vMerge value on the raw <w:tc> at (row_idx, col_idx).

    Goes via raw XML because python-docx's ``Table.cell(r, c)`` and
    ``Row.cells[c]`` skip across vMerge continuation cells and return the
    anchor cell — useful for end-user APIs but masks the underlying merge
    structure we want to verify here.
    """
    from docx.oxml.ns import qn
    trs = table._element.findall(qn("w:tr"))
    if not (0 <= row_idx < len(trs)):
        return ("MISSING_ROW", None)
    tcs = trs[row_idx].findall(qn("w:tc"))
    if not (0 <= col_idx < len(tcs)):
        return ("MISSING_TC", None)
    tcPr = tcs[col_idx].find(qn("w:tcPr"))
    if tcPr is None:
        return None
    vm = tcPr.find(qn("w:vMerge"))
    if vm is None:
        return None
    return vm.get(qn("w:val")) or ""

def test_add_table_auto_merges_empty_leading_cells_in_subrows(workdir):
    """Mimic the LLM pattern: header row + multiple sub-rows where leading
    columns are blank for sub-entries. With auto_merge_empty=True (default)
    the blanks should become OOXML vMerge continuations of the cell above —
    rendered as a single visually merged cell in PDF.
    """
    src = "in.docx"
    Document().save(workdir / src)

    out = editor.apply_edits(
        input_filename=src, output_filename="t.docx",
        ops=[{
            "op": "add_table",
            "position": "end",
            "has_header": True,
            "rows": [
                ["序号", "建议方向",  "重点任务",     "关键举措"],
                ["一",  "夯实基础",  "算力基础设施", "...A"],
                ["",    "",          "数据要素",     "...B"],
                ["",    "",          "基础软件",     "...C"],
                ["二",  "强化创新",  "关键核心技术", "...D"],
                ["",    "",          "培育企业",     "...E"],
            ],
        }],
    )
    assert out["ops_succeeded"] == 1
    res = out["results"][0]
    # 4 vmerge groups: col0 row1-3 ("一"/empty/empty), col1 row1-3, col0 row4-5, col1 row4-5
    assert res["vertical_merge_groups"] == 4

    d = Document(workdir / "t.docx")
    tbl = d.tables[-1]
    # Header row (idx 0) gets no vMerge
    assert _vmerge_at(tbl, 0, 0) is None
    # Group 1, col 0:  row1 = restart, row2 = continue, row3 = continue
    assert _vmerge_at(tbl, 1, 0) == "restart"
    assert _vmerge_at(tbl, 2, 0) == ""
    assert _vmerge_at(tbl, 3, 0) == ""
    # Row 4 starts a new group → restart
    assert _vmerge_at(tbl, 4, 0) == "restart"
    assert _vmerge_at(tbl, 5, 0) == ""
    # Same for col 1
    assert _vmerge_at(tbl, 1, 1) == "restart"
    assert _vmerge_at(tbl, 2, 1) == ""
    assert _vmerge_at(tbl, 3, 1) == ""
    # Cols 2 + 3 have no consecutive blanks → no vMerge
    assert _vmerge_at(tbl, 1, 2) is None
    assert _vmerge_at(tbl, 1, 3) is None

def test_add_table_no_auto_merge_when_disabled(workdir):
    src = "in.docx"
    Document().save(workdir / src)
    out = editor.apply_edits(
        input_filename=src, output_filename="t.docx",
        ops=[{
            "op": "add_table",
            "position": "end",
            "has_header": True,
            "auto_merge_empty": False,
            "rows": [
                ["A", "B"],
                ["1", "2"],
                ["",  ""],
                ["",  ""],
            ],
        }],
    )
    assert out["results"][0]["vertical_merge_groups"] == 0
    d = Document(workdir / "t.docx")
    tbl = d.tables[-1]
    for r in range(4):
        for c in range(2):
            assert _vmerge_at(tbl, r, c) is None

# ── style applies to FIRST paragraph only by default (regression) ──────────

def _para_style(p):
    return p.style.name if p.style else ""

def test_insert_text_style_applies_only_to_first_line(workdir):
    """Regression for chat_20260509_105809: model passed style="Heading 3"
    expecting only the first line to be a heading and the rest to be body
    paragraphs. Old behaviour: every \\n-split line got Heading 3 + heading
    font, turning all body text into headings.
    """
    src = "in.docx"
    Document().save(workdir / src)
    editor.apply_edits(
        input_filename=src, output_filename="t.docx",
        ops=[{
            "op": "insert",
            "text": "3.3 特色企业分析\n\n位于长三角南翼，是制造业重镇。\n\n在智能制造领域，拥有多家代表性企业。",
            "style": "Heading 3",
            "position": "end",
        }],
    )
    d = Document(workdir / "t.docx")
    inserted = [p for p in d.paragraphs if p.text]
    # Find the inserted lines.
    h = next(p for p in inserted if p.text.startswith("3.3"))
    body1 = next(p for p in inserted if "位于" in p.text)
    body2 = next(p for p in inserted if "在智能制造领域" in p.text)
    assert _para_style(h) == "Heading 3"
    # Body lines must NOT inherit the Heading 3 style.
    assert _para_style(body1) != "Heading 3"
    assert _para_style(body2) != "Heading 3"
    assert _para_style(body1) in ("Normal", "")
    assert _para_style(body2) in ("Normal", "")

def test_insert_text_style_for_all_when_explicit_opt_in(workdir):
    """Opt-in: style_for_all=True applies the style to every line (e.g. lists)."""
    src = "in.docx"
    Document().save(workdir / src)
    editor.apply_edits(
        input_filename=src, output_filename="t.docx",
        ops=[{
            "op": "insert",
            "text": "Line A\nLine B\nLine C",
            "style": "Heading 3",
            "style_for_all": True,
            "position": "end",
        }],
    )
    d = Document(workdir / "t.docx")
    matched = [p for p in d.paragraphs if p.text.startswith("Line ")]
    assert len(matched) == 3
    for p in matched:
        assert _para_style(p) == "Heading 3"

# ── format="markdown" — render headings / bullets / tables / inline ────────

def _block_kinds_after(d: Document, anchor_text: str) -> list[tuple[str, str]]:
    """Return (kind, text) tuples for body-level children appearing AFTER the
    paragraph whose text contains ``anchor_text``, until next heading at same
    or higher level."""
    from docx.oxml.ns import qn
    body = d.element.body
    seen = False
    out = []
    for child in body:
        if child.tag == qn("w:p"):
            txt = "".join(t.text or "" for t in child.iter(qn("w:t")))
            if anchor_text in txt:
                seen = True
                continue
            if seen:
                # Detect a heading via pStyle
                pstyle = child.find(f"{qn('w:pPr')}/{qn('w:pStyle')}")
                style_val = pstyle.get(qn("w:val")) if pstyle is not None else ""
                out.append(("p", txt, style_val))
        elif child.tag == qn("w:tbl") and seen:
            first_cell = child.find(f"{qn('w:tr')}/{qn('w:tc')}")
            label = "".join(t.text or "" for t in first_cell.iter(qn("w:t"))) if first_cell is not None else ""
            out.append(("tbl", label, ""))
    return out

def test_insert_markdown_renders_headings_and_lists(workdir):
    """format='markdown' renders ATX headings / bullets / numbered lists as
    proper Word block styles instead of literal text.
    """
    src = "in.docx"
    d0 = Document()
    d0.add_heading("第三章 企业画像与区域分布", level=2)
    d0.add_heading("3.2 区域分布格局", level=3)
    d0.add_paragraph("3.2 正文")
    d0.save(workdir / src)

    md = (
        "### 3.3 当地特色企业\n\n"
        "作为浙江省人工智能产业发展的重要城市之一。\n\n"
        "#### 3.3.1 龙头企业\n\n"
        "拥有一批 **龙头企业**，主要包括：\n\n"
        "- 舜宇光学\n"
        "- 均普智能\n"
        "- 奥克斯集团\n"
    )
    out = editor.apply_edits(
        input_filename=src, output_filename="t.docx",
        ops=[{
            "op": "insert",
            "text": md,
            "format": "markdown",
            "position": "after_heading",
            "anchor": "3.2 区域分布格局",
        }],
    )
    assert out["ops_succeeded"] == 1
    res = out["results"][0]
    assert res["format"] == "markdown"
    assert res["inserted_paragraphs"] >= 6  # 1 H3 + 1 body + 1 H4 + 1 body w/ bold + 3 bullets

    d = Document(workdir / "t.docx")

    # H3 + H4 must be PROPER heading styles (not literal "###" text).
    h3 = next((p for p in d.paragraphs if p.text.strip() == "3.3 当地特色企业"), None)
    h4 = next((p for p in d.paragraphs if p.text.strip() == "3.3.1 龙头企业"), None)
    assert h3 is not None and h3.style.name == "Heading 3"
    assert h4 is not None and h4.style.name == "Heading 4"

    # No literal "###" / "####" should appear.
    full_text = "\n".join(p.text for p in d.paragraphs)
    assert "###" not in full_text
    assert "####" not in full_text

    # Bullet list paragraphs should use List Bullet style.
    bullet_styles = [p.style.name for p in d.paragraphs if p.text in ("舜宇光学", "均普智能", "奥克斯集团")]
    assert all(s == "List Bullet" for s in bullet_styles), bullet_styles

    # Inline bold: paragraph "拥有一批 龙头企业..." should have a bold run.
    bold_paragraph = next((p for p in d.paragraphs if "龙头企业" in p.text and "" in p.text), None)
    assert bold_paragraph is not None
    assert any(r.bold for r in bold_paragraph.runs if r.text == "龙头企业")

def test_insert_markdown_rejects_embedded_table(workdir):
    """Tables MUST go through word_add_table; embedding ``| a | b |`` rows in
    markdown text is rejected with a clear redirect. Keeps the boundary
    between text-ops and table-ops crisp and prevents caption duplication.
    """
    src = "in.docx"
    d0 = Document()
    d0.add_heading("3.3 标题", level=3)
    d0.add_paragraph("正文")
    d0.save(workdir / src)

    md = (
        "新增表格：\n\n"
        "| 企业 | 领域 |\n"
        "| --- | --- |\n"
        "| 舜宇光学 | 光学 |\n"
    )
    out = editor.apply_edits(
        input_filename=src, output_filename="t.docx",
        ops=[{
            "op": "insert",
            "text": md,
            "format": "markdown",
            "position": "after_heading",
            "anchor": "3.3 标题",
        }],
    )
    assert out["ops_failed"] == 1
    err = out["results"][0]["error"]
    assert "table" in err.lower()
    assert "word_add_table" in err

# ── Caption deduplication when add_table follows insert_text(caption) ─────

def test_add_table_caption_suppressed_when_anchor_paragraph_already_has_caption(workdir):
    """Regression for chat_20260509_121616: model wrote a caption paragraph
    via insert_text("**表3-5 ...**"), then called add_table with the SAME
    string as both caption= and anchor=. Old behaviour produced two visible
    captions stacked; the new behaviour detects the duplicate and emits
    only the table.
    """
    src = "in.docx"
    d0 = Document()
    d0.add_paragraph("前言")
    d0.add_paragraph("表3-5 人工智能产业链代表性企业")  # the caption already exists
    d0.add_paragraph("尾段")
    d0.save(workdir / src)

    out = editor.apply_edits(
        input_filename=src, output_filename="t.docx",
        ops=[{
            "op": "add_table",
            "rows": [["企业", "领域"], ["舜宇", "光学"]],
            "caption": "表3-5 人工智能产业链代表性企业",
            "anchor":  "表3-5 人工智能产业链代表性企业",
            "position": "after_paragraph",
        }],
    )
    assert out["ops_succeeded"] == 1
    res = out["results"][0]
    assert res["caption_suppressed_duplicate"] is True
    assert res["caption"] is None  # caption was wiped, not emitted

    d = Document(workdir / "t.docx")
    # Count caption paragraphs — must be exactly 1 (the original).
    captions = [p for p in d.paragraphs if p.text.strip() == "表3-5 人工智能产业链代表性企业"]
    assert len(captions) == 1, [p.text for p in d.paragraphs]

def test_add_table_caption_kept_when_no_duplicate(workdir):
    """Sanity: when caption is genuinely new (not equal to the anchor para),
    it IS emitted as before."""
    src = "in.docx"
    d0 = Document()
    d0.add_paragraph("Some intro")
    d0.add_paragraph("anchor here")
    d0.add_paragraph("trailing")
    d0.save(workdir / src)

    out = editor.apply_edits(
        input_filename=src, output_filename="t.docx",
        ops=[{
            "op": "add_table",
            "rows": [["a", "b"], ["1", "2"]],
            "caption": "Brand new caption",
            "anchor":  "anchor here",
            "position": "after_paragraph",
        }],
    )
    assert out["results"][0]["caption_suppressed_duplicate"] is False
    d = Document(workdir / "t.docx")
    assert any(p.text == "Brand new caption" for p in d.paragraphs)

# ── after_paragraph skips tables that immediately follow the anchor ────────

def test_after_paragraph_skips_following_tables(workdir):
    """When anchor paragraph is directly followed by a table (the typical
    'caption + its data table' shape), inserting "after the caption" should
    land AFTER the table, not in the gap between them. This was the
    regression that wedged 3.3 between '表3-4 caption' and 表3-4's data.
    """
    src = "in.docx"
    d0 = Document()
    d0.add_paragraph("表3-4 城市分布TOP5")  # caption-like anchor
    t = d0.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "city"; t.cell(0, 1).text = "n"
    t.cell(1, 0).text = "Beijing"; t.cell(1, 1).text = "100"
    d0.add_paragraph("post-table tail")
    d0.save(workdir / src)

    out = editor.apply_edits(
        input_filename=src, output_filename="t.docx",
        ops=[{
            "op": "insert",
            "text": "INSERTED LINE",
            "anchor": "表3-4 城市分布TOP5",
            "position": "after_paragraph",
        }],
    )
    assert out["ops_succeeded"] == 1
    d = Document(workdir / "t.docx")
    # Walk body in order and check INSERTED LINE comes AFTER the <w:tbl>.
    from docx.oxml.ns import qn
    body = d.element.body
    seq = []
    for c in body:
        if c.tag == qn("w:p"):
            txt = "".join(t.text or "" for t in c.iter(qn("w:t")))
            seq.append(("p", txt))
        elif c.tag == qn("w:tbl"):
            seq.append(("tbl", "table"))
    # Required order: caption, tbl, INSERTED LINE, post-table tail
    cap_idx = next(i for i, (k, t) in enumerate(seq) if "表3-4" in t)
    ins_idx = next(i for i, (k, t) in enumerate(seq) if t == "INSERTED LINE")
    tbl_idx = next(i for i, (k, t) in enumerate(seq) if k == "tbl")
    tail_idx = next(i for i, (k, t) in enumerate(seq) if t == "post-table tail")
    assert cap_idx < tbl_idx < ins_idx < tail_idx, seq

def test_after_paragraph_no_following_table_unchanged(workdir):
    """Sanity: when anchor has no following table, after_paragraph behaves
    as a literal insert-immediately-after."""
    src = "in.docx"
    d0 = Document()
    d0.add_paragraph("para A")
    d0.add_paragraph("para B")
    d0.save(workdir / src)
    editor.apply_edits(
        input_filename=src, output_filename="t.docx",
        ops=[{"op": "insert", "text": "X", "anchor": "para A",
              "position": "after_paragraph"}],
    )
    d = Document(workdir / "t.docx")
    texts = [p.text for p in d.paragraphs]
    assert texts == ["para A", "X", "para B"]

def test_after_paragraph_skips_multiple_following_tables(workdir):
    """If there are MULTIPLE tables (rare but possible) directly after the
    anchor, all of them are skipped — the insertion lands after the last."""
    src = "in.docx"
    d0 = Document()
    d0.add_paragraph("anchor")
    for i in range(3):
        t = d0.add_table(rows=1, cols=1)
        t.cell(0, 0).text = f"T{i}"
    d0.add_paragraph("post")
    d0.save(workdir / src)
    editor.apply_edits(
        input_filename=src, output_filename="t.docx",
        ops=[{"op": "insert", "text": "X", "anchor": "anchor",
              "position": "after_paragraph"}],
    )
    d = Document(workdir / "t.docx")
    from docx.oxml.ns import qn
    seq = []
    for c in d.element.body:
        if c.tag == qn("w:p"):
            seq.append(("p", "".join(t.text or "" for t in c.iter(qn("w:t")))))
        elif c.tag == qn("w:tbl"):
            seq.append(("tbl", c.find(f"{qn('w:tr')}/{qn('w:tc')}").iter(qn("w:t")).__next__().text))
    # Order: anchor, T0, T1, T2, X, post
    assert seq[0] == ("p", "anchor")
    assert seq[1][0] == "tbl" and seq[2][0] == "tbl" and seq[3][0] == "tbl"
    assert seq[4] == ("p", "X")
    assert seq[5] == ("p", "post")

# ── after_section: insert at the END of a heading section ─────────────────

def test_after_section_inserts_before_next_sibling_heading(workdir):
    """Regression for chat_20260509_123423: model wanted to add 3.3 after
    the 3.2 section. anchoring on a caption inside 3.2 + after_paragraph
    misplaced 3.3 mid-section, leaving the 3.2 trailing paragraphs stranded
    INSIDE the new 3.3. position="after_section" handles this cleanly: it
    walks from the anchor heading to the next same-or-higher-level heading
    and inserts before it.
    """
    src = "in.docx"
    d0 = Document()
    d0.add_heading("3.2 区域分布格局", level=3)
    d0.add_paragraph("3.2 引言")
    d0.add_paragraph("表3-3 caption")
    t1 = d0.add_table(rows=2, cols=1)
    t1.cell(0, 0).text = "h1"; t1.cell(1, 0).text = "v1"
    d0.add_paragraph("表3-4 caption")
    t2 = d0.add_table(rows=2, cols=1)
    t2.cell(0, 0).text = "h2"; t2.cell(1, 0).text = "v2"
    d0.add_paragraph("北京市以3,098家企业...")  # 3.2 trailing body
    d0.add_paragraph("深圳市以2,123家企业...")  # 3.2 trailing body
    d0.add_heading("第四章 技术创新", level=2)
    d0.save(workdir / src)

    out = editor.apply_edits(
        input_filename=src, output_filename="t.docx",
        ops=[{
            "op": "insert",
            "format": "markdown",
            "position": "after_section",
            "anchor": "3.2 区域分布格局",
            "text": "### 3.3 特色企业\n\n概述...\n\n#### 子标题",
        }],
    )
    assert out["ops_succeeded"] == 1, out

    d = Document(workdir / "t.docx")
    from docx.oxml.ns import qn as _qn
    seq = []
    for c in d.element.body:
        if c.tag == _qn("w:p"):
            txt = "".join(t.text or "" for t in c.iter(_qn("w:t")))
            seq.append(("p", txt))
        elif c.tag == _qn("w:tbl"):
            seq.append(("tbl", "table"))
    # The new 3.3 must come AFTER all of 3.2's content (incl. trailing
    # paragraphs and tables) and BEFORE the 第四章 heading.
    bj_idx   = next(i for i, (k, t) in enumerate(seq) if "北京市以3,098家" in t)
    sz_idx   = next(i for i, (k, t) in enumerate(seq) if "深圳市以2,123家" in t)
    h33_idx  = next(i for i, (k, t) in enumerate(seq) if t == "3.3 特色企业")
    chap4_idx = next(i for i, (k, t) in enumerate(seq) if t == "第四章 技术创新")
    assert bj_idx < h33_idx and sz_idx < h33_idx, "3.2 trailing body must stay in 3.2"
    assert h33_idx < chap4_idx, "new 3.3 must precede 第四章"

def test_after_section_falls_back_to_doc_end_when_no_next_heading(workdir):
    """If the anchored heading is the LAST same-level heading in the doc,
    after_section should append at end of body."""
    src = "in.docx"
    d0 = Document()
    d0.add_heading("结论", level=2)
    d0.add_paragraph("conclusion body 1")
    d0.add_paragraph("conclusion body 2")
    d0.save(workdir / src)

    editor.apply_edits(
        input_filename=src, output_filename="t.docx",
        ops=[{"op": "insert", "position": "after_section",
              "anchor": "结论", "text": "appended"}],
    )
    d = Document(workdir / "t.docx")
    texts = [p.text for p in d.paragraphs]
    assert texts[-1] == "appended"

def test_after_section_requires_anchor(workdir):
    src = "in.docx"
    Document().save(workdir / src)
    out = editor.apply_edits(
        input_filename=src, output_filename="t.docx",
        ops=[{"op": "insert", "position": "after_section", "text": "x"}],
    )
    assert out["ops_failed"] == 1
    assert "anchor" in out["results"][0]["error"]

def test_add_table_after_section(workdir):
    """add_table now supports position=after_section so apply_edits batches
    that mix insert + add_table can use one position semantics for both
    (regression for chat_20260509_124517 where add_table rejected
    after_section, breaking the cascade).
    """
    src = "in.docx"
    d0 = Document()
    d0.add_heading("2.4 下游应用层分析", level=3)
    d0.add_paragraph("2.4 引言")
    d0.add_paragraph("2.4 后续")
    d0.add_heading("第三章", level=2)
    d0.save(workdir / src)

    out = editor.apply_edits(
        input_filename=src, output_filename="t.docx",
        ops=[{
            "op": "add_table",
            "rows": [["A", "B"], ["1", "2"]],
            "position": "after_section",
            "anchor": "2.4 下游应用层分析",
            "caption": "表X 测试",
        }],
    )
    assert out["ops_succeeded"] == 1, out["results"]
    d = Document(workdir / "t.docx")
    from docx.oxml.ns import qn as _qn
    seq = []
    for c in d.element.body:
        if c.tag == _qn("w:p"):
            seq.append(("p", "".join(t.text or "" for t in c.iter(_qn("w:t")))))
        elif c.tag == _qn("w:tbl"):
            seq.append(("tbl", "table"))
    # Order must be: 2.4 heading, 2.4 引言, 2.4 后续, caption, table, 第三章
    cap_idx = next(i for i, (k, t) in enumerate(seq) if t == "表X 测试")
    tbl_idx = next(i for i, (k, t) in enumerate(seq) if k == "tbl")
    chap3_idx = next(i for i, (k, t) in enumerate(seq) if t == "第三章")
    after_idx = next(i for i, (k, t) in enumerate(seq) if t == "2.4 后续")
    assert after_idx < cap_idx < tbl_idx < chap3_idx, seq

def test_add_table_after_section_falls_back_to_doc_end(workdir):
    """When the anchored heading is the LAST in the doc, after_section
    appends at body end (before sectPr)."""
    src = "in.docx"
    d0 = Document()
    d0.add_heading("结论", level=2)
    d0.add_paragraph("结论正文")
    d0.save(workdir / src)
    out = editor.apply_edits(
        input_filename=src, output_filename="t.docx",
        ops=[{"op": "add_table", "rows": [["a"], ["b"]],
              "position": "after_section", "anchor": "结论"}],
    )
    assert out["ops_succeeded"] == 1
    d = Document(workdir / "t.docx")
    # The new table must come AFTER the conclusion body.
    from docx.oxml.ns import qn as _qn
    seq = []
    for c in d.element.body:
        if c.tag == _qn("w:p"):
            seq.append(("p", "".join(t.text or "" for t in c.iter(_qn("w:t")))))
        elif c.tag == _qn("w:tbl"):
            seq.append(("tbl", "tbl"))
    body_idx = next(i for i, (k, t) in enumerate(seq) if t == "结论正文")
    tbl_idx = next(i for i, (k, t) in enumerate(seq) if k == "tbl")
    assert body_idx < tbl_idx

# ── Auto-redirect: after_heading + sibling-heading text → after_section ────

def test_insert_markdown_after_heading_with_sibling_heading_redirects(workdir):
    """Regression for chat_20260509_124517: model called insert with
    position='after_heading', anchor='2.4 ...', text='### 2.5 ...' (markdown
    heading at SAME level as anchor — clearly a new sibling section). Naive
    after_heading wedges 2.5 between 2.4 heading and 2.4 body. The auto-
    redirect converts it to after_section so 2.5 lands at the end of 2.4's
    section.
    """
    src = "in.docx"
    d0 = Document()
    d0.add_heading("2.4 下游应用层分析", level=3)
    d0.add_paragraph("2.4 引言段")
    t = d0.add_table(rows=1, cols=1)
    t.cell(0, 0).text = "data"
    d0.add_paragraph("2.4 后续正文")
    d0.add_heading("第三章", level=2)
    d0.save(workdir / src)

    out = editor.apply_edits(
        input_filename=src, output_filename="t.docx",
        ops=[{
            "op": "insert",
            "format": "markdown",
            "position": "after_heading",
            "anchor": "2.4 下游应用层分析",
            "text": "### 2.5 特色企业\n\n2.5 引言段...",
        }],
    )
    assert out["ops_succeeded"] == 1
    res = out["results"][0]
    # Result tells the LLM the redirect happened.
    assert res.get("position") == "after_section"
    assert res.get("position_redirected_from") == "after_heading"

    d = Document(workdir / "t.docx")
    from docx.oxml.ns import qn as _qn
    seq = []
    for c in d.element.body:
        if c.tag == _qn("w:p"):
            seq.append(("p", "".join(t.text or "" for t in c.iter(_qn("w:t")))))
        elif c.tag == _qn("w:tbl"):
            seq.append(("tbl", "tbl"))
    h24_idx = next(i for i, (k, t) in enumerate(seq) if t == "2.4 下游应用层分析")
    intro_idx = next(i for i, (k, t) in enumerate(seq) if t == "2.4 引言段")
    tbl_idx = next(i for i, (k, t) in enumerate(seq) if k == "tbl")
    tail_idx = next(i for i, (k, t) in enumerate(seq) if t == "2.4 后续正文")
    h25_idx = next(i for i, (k, t) in enumerate(seq) if t == "2.5 特色企业")
    chap3_idx = next(i for i, (k, t) in enumerate(seq) if t == "第三章")
    # 2.5 must come AFTER all of 2.4's content and BEFORE chapter 3.
    assert h24_idx < intro_idx < tbl_idx < tail_idx < h25_idx < chap3_idx, seq

def test_insert_markdown_after_heading_no_redirect_when_text_is_body(workdir):
    """When the text is plain body content (no leading heading marker), the
    user really DOES want it right under the anchor heading — no redirect.
    """
    src = "in.docx"
    d0 = Document()
    d0.add_heading("2.4 下游应用层", level=3)
    d0.add_paragraph("existing body")
    d0.add_heading("第三章", level=2)
    d0.save(workdir / src)

    out = editor.apply_edits(
        input_filename=src, output_filename="t.docx",
        ops=[{
            "op": "insert",
            "format": "markdown",
            "position": "after_heading",
            "anchor": "2.4 下游应用层",
            "text": "本节将探讨...",  # plain body, no heading marker
        }],
    )
    res = out["results"][0]
    assert res["ok"] is True
    assert res.get("position") == "after_heading"
    assert "position_redirected_from" not in res

def test_insert_markdown_after_heading_lower_level_text_no_redirect(workdir):
    """When the text starts with a heading but at a LOWER level (e.g.
    inserting a 2.4.1 sub-heading under 2.4), no redirect — that's
    legitimate after_heading usage (adding a sub-heading right under)."""
    src = "in.docx"
    d0 = Document()
    d0.add_heading("2.4 下游应用层", level=3)
    d0.add_paragraph("existing body")
    d0.add_heading("第三章", level=2)
    d0.save(workdir / src)

    out = editor.apply_edits(
        input_filename=src, output_filename="t.docx",
        ops=[{
            "op": "insert",
            "format": "markdown",
            "position": "after_heading",
            "anchor": "2.4 下游应用层",
            "text": "#### 2.4.1 子节\n\n子节内容",  # H4 < H3 anchor → no redirect
        }],
    )
    res = out["results"][0]
    assert res["ok"] is True
    assert res.get("position") == "after_heading"
    assert "position_redirected_from" not in res

def test_after_section_unknown_anchor_fails_cleanly(workdir):
    src = "in.docx"
    d0 = Document()
    d0.add_heading("Real Heading", level=2)
    d0.save(workdir / src)
    out = editor.apply_edits(
        input_filename=src, output_filename="t.docx",
        ops=[{"op": "insert", "position": "after_section",
              "anchor": "Nonexistent Heading", "text": "x"}],
    )
    assert out["ops_failed"] == 1
    assert "no heading containing" in out["results"][0]["error"]

def test_after_paragraph_markdown_insert_skips_following_table(workdir):
    """Same skip-tables behaviour applies to format='markdown' inserts."""
    src = "in.docx"
    d0 = Document()
    d0.add_paragraph("表3-4 caption")
    t = d0.add_table(rows=1, cols=1)
    t.cell(0, 0).text = "data"
    d0.add_paragraph("tail")
    d0.save(workdir / src)
    editor.apply_edits(
        input_filename=src, output_filename="t.docx",
        ops=[{
            "op": "insert",
            "text": "### NEW SECTION\n\nbody text",
            "format": "markdown",
            "anchor": "表3-4 caption",
            "position": "after_paragraph",
        }],
    )
    d = Document(workdir / "t.docx")
    from docx.oxml.ns import qn
    seq = []
    for c in d.element.body:
        if c.tag == qn("w:p"):
            seq.append(("p", "".join(t.text or "" for t in c.iter(qn("w:t")))))
        elif c.tag == qn("w:tbl"):
            seq.append(("tbl", "table"))
    cap_idx = next(i for i, (k, t) in enumerate(seq) if "表3-4 caption" in t)
    tbl_idx = next(i for i, (k, t) in enumerate(seq) if k == "tbl")
    new_section_idx = next(i for i, (k, t) in enumerate(seq) if t == "NEW SECTION")
    tail_idx = next(i for i, (k, t) in enumerate(seq) if t == "tail")
    assert cap_idx < tbl_idx < new_section_idx < tail_idx

def test_insert_text_warns_when_forced_text_on_markdown(workdir):
    """Explicit format='text' on markdown-looking input → warning + literal."""
    src = "in.docx"
    Document().save(workdir / src)
    out = editor.apply_edits(
        input_filename=src, output_filename="t.docx",
        ops=[{
            "op": "insert",
            "text": "### 这是一个 markdown 标题\n\n**这是粗体**\n\n- 项目一\n- 项目二",
            "position": "end",
            "format": "text",  # force literal — opt out of auto-rendering
        }],
    )
    assert out["ops_succeeded"] == 1
    res = out["results"][0]
    assert res["format"] == "text"
    assert "warning" in res
    assert "markdown" in res["warning"].lower()
    # literal: the '### ...' markup survives verbatim as paragraph text
    d = Document(workdir / "t.docx")
    assert any("### 这是一个 markdown 标题" in p.text for p in d.paragraphs)


def test_insert_text_auto_renders_markdown_by_default(workdir):
    """No format → 'auto' renders markdown signals as real Word blocks."""
    src = "in.docx"
    Document().save(workdir / src)
    out = editor.apply_edits(
        input_filename=src, output_filename="t.docx",
        ops=[{
            "op": "insert",
            "text": "### 这是一个 markdown 标题\n\n**这是粗体**\n\n- 项目一\n- 项目二",
            "position": "end",
            # NOTE: no format → defaults to "auto" → renders as markdown
        }],
    )
    assert out["ops_succeeded"] == 1
    res = out["results"][0]
    assert res["format"] == "markdown"
    assert "warning" not in res
    d = Document(workdir / "t.docx")
    # heading rendered as a real Heading style, not literal '### ...'
    assert any(
        "markdown 标题" in p.text and (p.style.name or "").startswith("Heading")
        for p in d.paragraphs
    )
    assert not any("### " in p.text for p in d.paragraphs)

def test_insert_markdown_block_relocated_to_anchor_in_order(workdir):
    """The newly emitted blocks land in the correct spatial position (right
    after the anchor heading), in document order.
    """
    src = "in.docx"
    d0 = Document()
    d0.add_heading("A", level=1)
    d0.add_paragraph("alpha body")
    d0.add_heading("B", level=1)
    d0.add_paragraph("beta body")
    d0.save(workdir / src)

    md = "### inserted h3\n\nfirst inserted body\n\nsecond inserted body\n"
    editor.apply_edits(
        input_filename=src, output_filename="t.docx",
        ops=[{
            "op": "insert",
            "text": md,
            "format": "markdown",
            "position": "after_heading",
            "anchor": "A",
        }],
    )
    d = Document(workdir / "t.docx")
    para_texts = [p.text for p in d.paragraphs if p.text.strip()]
    # Original "B" heading must come AFTER all inserted content and beta_body.
    a_idx = para_texts.index("A")
    h3_idx = para_texts.index("inserted h3")
    f_idx = para_texts.index("first inserted body")
    b_idx = para_texts.index("B")
    assert a_idx < h3_idx < f_idx < b_idx

def test_insert_text_no_style_keeps_normal_for_all(workdir):
    """Sanity: omitting style leaves all lines as default (Normal)."""
    src = "in.docx"
    Document().save(workdir / src)
    editor.apply_edits(
        input_filename=src, output_filename="t.docx",
        ops=[{
            "op": "insert",
            "text": "Line A\nLine B",
            "position": "end",
        }],
    )
    d = Document(workdir / "t.docx")
    for p in [p for p in d.paragraphs if p.text.startswith("Line ")]:
        assert _para_style(p) in ("Normal", "")

# ── word_add_table position is required (regression) ──────────────────────

@pytest.mark.asyncio
async def test_apply_edits_add_table_op_requires_position(monkeypatch, workdir, src_doc):
    """The apply_edits ``add_table`` op also requires explicit position — a
    missing key raises TypeError captured per-op, so the batch surfaces a
    clean failure (not a silent end-of-doc dump)."""
    out = editor.apply_edits(
        input_filename=src_doc, output_filename="t.docx",
        ops=[{"op": "add_table", "rows": [["A", "B"], ["1", "2"]]}],  # no position
    )
    assert out["ops_failed"] == 1
    assert out["ops_succeeded"] == 0
    err = out["results"][0]["error"]
    assert "position" in err

def test_add_table_blanks_under_blanks_dont_extend_into_new_group(workdir):
    """When a column has a non-blank → blank → non-blank pattern, the blank
    only merges with the cell above, NOT downward.
    """
    src = "in.docx"
    Document().save(workdir / src)
    out = editor.apply_edits(
        input_filename=src, output_filename="t.docx",
        ops=[{
            "op": "add_table",
            "position": "end",
            "has_header": True,
            "rows": [
                ["A", "B"],
                ["x", "y"],   # row 1
                ["",  "z"],   # row 2  — col 0 blank → merges UP into "x"
                ["w", "u"],   # row 3  — col 0 "w" starts fresh, NOT a continuation
            ],
        }],
    )
    # One merge group: col 0 rows 1-2
    assert out["results"][0]["vertical_merge_groups"] == 1
    d = Document(workdir / "t.docx")
    tbl = d.tables[-1]
    assert _vmerge_at(tbl, 1, 0) == "restart"
    assert _vmerge_at(tbl, 2, 0) == ""
    # Row 3 col 0 must NOT be marked vMerge — it's a new value, not a continuation.
    assert _vmerge_at(tbl, 3, 0) is None
    # Col 1 has no blanks → no vMerge anywhere.
    for r in range(1, 4):
        assert _vmerge_at(tbl, r, 1) is None
