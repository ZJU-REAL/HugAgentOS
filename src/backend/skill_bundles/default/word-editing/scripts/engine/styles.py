"""CJK font and heading style helpers for python-docx.

Extracted from ``mcp_servers/report_export_mcp/impl.py:46-130`` so the same font
logic is shared between the legacy markdown-export tool and the new office_mcp
word tools. Behavior must remain identical to keep existing exports rendering
the same.
"""
from __future__ import annotations

from typing import Any

# 方正字体系列（与仓库顶层 resources/fonts/ 中的文件名对齐）
BODY_FONT = "方正仿宋简体"
HEADING_FONT = "方正小标宋简体"
CODE_FONT = "Courier New"

# 英文文档时使用的兜底字体
EN_BODY_FONT = "Calibri"
EN_HEADING_FONT = "Calibri"

# Table styling — matches report_export_mcp so user-facing rendering is identical
TABLE_HEADER_COLOR = "366092"
TABLE_ALT_ROW_COLOR = "DCE6F1"
TABLE_BORDER_COLOR = "B8CCE4"

# Single source of truth for ``position`` values across editor.py / builder.py /
# the word-editing skill scripts. Without this, four sites had to keep tuples in sync.
INSERT_POSITIONS: tuple[str, ...] = (
    "end", "start",
    "after_heading", "after_section",
    "after_paragraph", "before_paragraph",
)
ANCHOR_REQUIRED_POSITIONS: frozenset[str] = frozenset({
    "after_heading", "after_section",
    "after_paragraph", "before_paragraph",
})


def font_for_style(style_name: str | None, body_font: str = BODY_FONT, heading_font: str = HEADING_FONT) -> str:
    """Pick the right CJK font for a paragraph style name."""
    if not style_name:
        return body_font
    return heading_font if style_name.startswith(("Title", "Heading")) else body_font


def fonts_for_language(language: str) -> tuple[str, str]:
    """Return (body_font, heading_font) for a given language code.

    ``zh`` (default) uses 方正 fonts; anything else falls back to Calibri so
    the document still renders cleanly when the 方正 family isn't installed
    on the reader's machine.
    """
    if (language or "zh").lower() == "zh":
        return BODY_FONT, HEADING_FONT
    return EN_BODY_FONT, EN_HEADING_FONT


def set_document_default_fonts(doc, ea_font: str) -> None:
    """Write docDefaults into styles.xml so every paragraph inherits the CJK font."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    styles_el = doc.styles.element
    docDefaults = styles_el.find(qn("w:docDefaults"))
    if docDefaults is None:
        docDefaults = OxmlElement("w:docDefaults")
        styles_el.insert(0, docDefaults)

    rPrDefault = docDefaults.find(qn("w:rPrDefault"))
    if rPrDefault is None:
        rPrDefault = OxmlElement("w:rPrDefault")
        docDefaults.append(rPrDefault)

    rPr = rPrDefault.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        rPrDefault.append(rPr)

    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)

    rFonts.set(qn("w:eastAsia"), ea_font)
    rFonts.set(qn("w:cs"), ea_font)


def apply_cjk_font_to_para(para, font_name: str) -> None:
    """Apply CJK font to all runs + paragraph mark in a paragraph."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def _set_rFonts(rPr_el: Any, name: str) -> None:
        rFonts = rPr_el.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr_el.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), name)
        rFonts.set(qn("w:ascii"), name)
        rFonts.set(qn("w:hAnsi"), name)

    for run in para.runs:
        run.font.name = font_name
        _set_rFonts(run._element.get_or_add_rPr(), font_name)

    pPr = para._element.get_or_add_pPr()
    pRpr = pPr.find(qn("w:rPr"))
    if pRpr is None:
        pRpr = OxmlElement("w:rPr")
        pPr.append(pRpr)
    _set_rFonts(pRpr, font_name)


def setup_heading_styles(doc, heading_font: str | None = None) -> None:
    """Set CJK font and force black color on heading style definitions.

    Covers Title / Subtitle / Heading 1-6 so no title or sub-heading inherits
    the blue theme color baked into python-docx's default template. Styles are
    matched by display name while iterating ``doc.styles`` rather than
    ``doc.styles[name]`` — pandoc-authored documents emit heading style names
    that don't round-trip through python-docx's built-in name lookup.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import RGBColor

    font = heading_font or HEADING_FONT

    targets = {"Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3",
               "Heading 4", "Heading 5", "Heading 6"}
    for style in doc.styles:
        if (style.name or "") not in targets:
            continue
        try:
            style.font.name = font
            rPr = style.element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), font)
            # Force black color, remove theme overrides
            style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            color_el = rPr.find(qn("w:color"))
            if color_el is not None:
                color_el.set(qn("w:val"), "000000")
                for attr in (qn("w:themeColor"), qn("w:themeShade"), qn("w:themeTint")):
                    if attr in color_el.attrib:
                        del color_el.attrib[attr]
        except (KeyError, ValueError):
            continue


def style_table(table, body_font: str = BODY_FONT) -> None:
    """Apply header background, alternating row shading, borders, CJK font, and center the table.

    Mirrors ``report_export_mcp/impl.py:_style_tables`` so tables created via
    ``word_add_table`` look identical to tables in ``export_report_to_docx``.
    """
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import RGBColor

    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    def _set_cell_shading(cell, color: str) -> None:
        tc_pr = cell._element.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), color)
        tc_pr.append(shading)

    def _set_cell_borders(cell) -> None:
        tc_pr = cell._element.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), TABLE_BORDER_COLOR)
            borders.append(el)
        tc_pr.append(borders)

    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            _set_cell_borders(cell)
            if row_idx == 0:
                _set_cell_shading(cell, TABLE_HEADER_COLOR)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif row_idx % 2 == 0:
                _set_cell_shading(cell, TABLE_ALT_ROW_COLOR)

            for para in cell.paragraphs:
                apply_cjk_font_to_para(para, body_font)
