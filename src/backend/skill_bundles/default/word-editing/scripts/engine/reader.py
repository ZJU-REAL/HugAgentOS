"""Read operations on .docx — text extraction, outline, table summary.

CRITICAL — index consistency
============================

All paragraph indices returned by this module refer to **body paragraphs only**
(``doc.paragraphs``), matching the addressing space used by every editing op
in ``word/editor.py``. Table cell content is NOT iterated as part of the body
paragraph stream; you address table content through ``table_index`` / ``row``
/ ``col`` (or via the dedicated table-aware ops) instead.

Why this matters
----------------
A previous version iterated body + table-cell paragraphs together for
``get_text`` / ``get_outline``, producing a paragraph_count that included
table cells (e.g. 600+ for a doc with 152 body paragraphs). The LLM then
addressed those indexes back to edit ops (``delete_paragraph``,
``delete_range``, …) which iterate body-only and reported "out of range" —
sending the model into a multi-pass re-read storm trying to figure out the
mismatch. Now every public API in this module emits the same body-paragraph
index space the editor consumes, so an index seen in ``get_text`` /
``get_outline`` is always valid as an anchor in edit ops.
"""
from __future__ import annotations

from typing import Any

from ._handle import input_path


def _iter_body_paragraphs(doc):
    """Yield body paragraphs only — same definition as ``editor._iter_body_paragraphs``.

    Table cell paragraphs are NOT included here; they have their own addressing
    via ``table_index`` / ``row`` / ``col`` in the edit ops. Headers / footers
    are also excluded (same reason — distinct addressing).
    """
    yield from doc.paragraphs


def _summarize_table(doc, table, table_index: int, body_paragraphs: list[Any]) -> dict[str, Any]:
    """Build a compact descriptor of one table for ``get_outline``.

    Includes the preceding body-paragraph index (for spatial reasoning),
    its text snippet (often a caption), the dimensions, and the first
    cell's text as a quick fingerprint.
    """
    tbl_el = table._element
    body = doc.element.body

    # Find the body-paragraph index that precedes this table in document order.
    # Tables sit at the body level (not inside paragraphs), so we walk
    # body-level children and count <w:p> elements seen before the <w:tbl>.
    from docx.oxml.ns import qn
    p_tag = qn("w:p")
    tbl_tag = qn("w:tbl")
    preceding_p_index = None
    body_p_count = 0
    for child in body:
        if child is tbl_el:
            preceding_p_index = body_p_count - 1 if body_p_count > 0 else None
            break
        if child.tag == p_tag:
            body_p_count += 1

    preceding_text = ""
    if preceding_p_index is not None and 0 <= preceding_p_index < len(body_paragraphs):
        preceding_text = body_paragraphs[preceding_p_index].text.strip()

    rows = list(table.rows)
    n_rows = len(rows)
    n_cols = len(rows[0].cells) if rows else 0
    first_cell_text = rows[0].cells[0].text.strip() if rows and rows[0].cells else ""

    return {
        "table_index": table_index,
        "preceding_paragraph_index": preceding_p_index,
        "preceding_paragraph_text": preceding_text[:80],
        "rows": n_rows,
        "cols": n_cols,
        "first_cell_text": first_cell_text[:80],
    }


def get_text(
    *,
    input_filename: str,
    paragraph_range: tuple[int, int] | None = None,
) -> dict[str, Any]:
    """Extract body text from a .docx.

    Args:
        input_filename: source .docx in sandbox cwd
        paragraph_range: optional ``(start, end)`` 0-based half-open slice over
            body paragraphs; None / omitted returns full body text.
            **Indexes match the body-paragraph index space used by edit ops** —
            an index seen here is valid as a paragraph anchor in
            ``replace_paragraph`` / ``delete_paragraph`` / ``delete_range``.

    Returns:
        ``{"paragraph_count": int,           # body paragraphs only
           "selected_range": [start, end],
           "text": str,                      # body paragraphs joined with \\n
           "table_count": int,               # number of tables in the doc
           "note": "table_text_excluded"}``  # advisory for the LLM

    Note:
        Table cell content is NOT included in ``text``. To read a table, use
        the table descriptors emitted by ``get_outline`` (which list every
        table with its index + dimensions) and address cells via the
        table-aware ops.
    """
    from docx import Document

    doc = Document(str(input_path(input_filename)))
    paragraphs = list(_iter_body_paragraphs(doc))
    total = len(paragraphs)

    if paragraph_range is None:
        selected = paragraphs
        sel_range = [0, total]
    else:
        start, end = paragraph_range
        if not (0 <= start <= end <= total):
            raise ValueError(
                f"paragraph_range {paragraph_range} invalid for document with "
                f"{total} body paragraphs (note: paragraph_count refers to body "
                "only, table cells are addressed separately via table_index)"
            )
        selected = paragraphs[start:end]
        sel_range = [start, end]

    return {
        "paragraph_count": total,
        "selected_range": sel_range,
        "text": "\n".join(p.text for p in selected),
        "table_count": len(doc.tables),
        "note": (
            "indexes refer to body paragraphs only (matches edit-op addressing). "
            "Table cell content is NOT included; see get_outline.tables for table "
            "descriptors and use table_index / set_cell_text / fill_table to edit."
        ),
    }


def get_outline(*, input_filename: str) -> dict[str, Any]:
    """Return heading tree + table descriptors of a .docx.

    Each heading entry::

        {"level": 0-6, "text": str, "anchor": str,
         "paragraph_index": int, "style": str}

    Each table entry::

        {"table_index": int,
         "preceding_paragraph_index": int|null,
         "preceding_paragraph_text": str,    # often a caption — first 80 chars
         "rows": int, "cols": int,
         "first_cell_text": str}             # first 80 chars of cell (0,0)

    - ``level=0`` means Title style; 1-6 are Heading 1 - Heading 6.
    - ``anchor`` = heading text — pass back to anchor-based edit ops
      (insert ``after_heading``, ``replace_section``, ``replace_paragraph``,
      ``delete_paragraph``, ``format`` with ``anchor=...``) for stability
      across edits.
    - ``paragraph_index`` is the body-paragraph index (matches edit-op
      addressing). Stable within a single call; may shift after insert/delete
      ops in the same batch — prefer ``anchor`` for multi-step workflows.
    - ``tables`` lists every ``<w:tbl>`` in document order with its
      ``preceding_paragraph_index`` so you can locate them spatially without
      having to scan ``get_text``.
    """
    from docx import Document

    doc = Document(str(input_path(input_filename)))
    paragraphs = list(_iter_body_paragraphs(doc))
    outline: list[dict[str, Any]] = []

    for idx, para in enumerate(paragraphs):
        style_name = para.style.name if para.style else ""
        # python-docx style names: "Title", "Heading 1", "Heading 2", ...
        level: int | None = None
        if style_name == "Title":
            level = 0  # treat title as level 0 (above H1)
        elif style_name.startswith("Heading "):
            try:
                level = int(style_name.split()[1])
            except (IndexError, ValueError):
                level = None
        if level is None:
            continue
        text = para.text.strip()
        if not text:
            continue
        outline.append(
            {
                "level": level,
                "text": text,
                "anchor": text,
                "paragraph_index": idx,
                "style": style_name,
            }
        )

    tables = [
        _summarize_table(doc, t, ti, paragraphs)
        for ti, t in enumerate(doc.tables)
    ]

    return {
        "heading_count": len(outline),
        "table_count": len(tables),
        "paragraph_count": len(paragraphs),
        "outline": outline,
        "tables": tables,
    }
