"""Markdown → python-docx block emission.

Used by ``editor._op_insert_text`` (when ``format="markdown"``) to render
Markdown text into proper Word block styles instead of literal characters.

Scope (deliberately narrow):
    - ATX headings ``#`` .. ``######`` → Heading 1-6
    - bullet lists ``- ``, ``* ``, ``+ `` → List Bullet
    - numbered lists ``1. ``, ``2. `` → List Number
    - fenced code blocks ``` ``` → Courier New paragraphs
    - inline ``**bold**`` / ``*italic*`` / `` `code` `` / ``[text](url)``

Out of scope:
    - Tables (``| a | b |``). Raised as ``ValueError`` so callers route to
      ``word_add_table`` which has explicit position / caption / merge
      controls. Mixing tables into markdown text led to caption duplication
      and merge-config losses.
    - Nested lists, blockquotes, images, raw HTML — not needed by the
      current LLM use cases.
"""
from __future__ import annotations

import re
from typing import Any


# Inline span matchers — mirrors ``report_export_mcp.impl._INLINE_RE``.
INLINE_RE = re.compile(
    r"(\*\*(.+?)\*\*)"        # 1,2 bold
    r"|(\*(.+?)\*)"           # 3,4 italic
    r"|(`(.+?)`)"             # 5,6 inline code
    r"|(\[(.+?)\]\((.+?)\))"  # 7,8,9 link
)

# Heuristic: any of these characters/sequences strongly suggest the input is
# Markdown rather than plain text. Used by ``insert_text format="text"`` to
# warn callers that their markdown markup will land as literal characters.
_SIGNAL_RE = re.compile(
    r"(?m)"  # ^ matches per-line
    r"(^#{1,6}\s)"        # ATX heading
    r"|(^[-*+]\s)"        # bullet
    r"|(^\d+\.\s)"        # numbered
    r"|(^\|.+\|\s*$)"     # table row
    r"|(```)"             # fenced code
    r"|(\*\*.+?\*\*)"     # bold span
)


def looks_like_markdown(text: str) -> bool:
    return bool(text and _SIGNAL_RE.search(text))


def add_inline_runs(paragraph, text: str, body_font: str, code_font: str = "Courier New") -> None:
    """Parse inline markdown (bold / italic / code / link) into runs."""
    from docx.shared import RGBColor

    last_end = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > last_end:
            run = paragraph.add_run(text[last_end:m.start()])
            run.font.name = body_font
        if m.group(2):
            run = paragraph.add_run(m.group(2))
            run.bold = True
            run.font.name = body_font
        elif m.group(4):
            run = paragraph.add_run(m.group(4))
            run.italic = True
            run.font.name = body_font
        elif m.group(6):
            run = paragraph.add_run(m.group(6))
            run.font.name = code_font
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif m.group(8):
            run = paragraph.add_run(m.group(8))
            run.font.name = body_font
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.underline = True
        last_end = m.end()
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.name = body_font


def emit_blocks(doc, markdown: str) -> list[Any]:
    """Append parsed markdown blocks to ``doc.element.body`` and return the
    OOXML elements that were appended (in document order). Caller can then
    relocate the returned slice via ``addnext`` / ``addprevious``.

    ``id()`` on lxml proxies is not stable across iterations, so the diff is
    done by counting block-element children before/after instead of set
    membership. The emit path only ever appends, so a count-based suffix
    take is safe.
    """
    from docx.oxml.ns import qn
    from .styles import (
        BODY_FONT,
        CODE_FONT,
        HEADING_FONT,
        apply_cjk_font_to_para,
    )

    body = doc.element.body
    p_tag, tbl_tag = qn("w:p"), qn("w:tbl")
    before_blocks = sum(1 for c in body if c.tag in (p_tag, tbl_tag))

    lines = (markdown or "").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")
        stripped = line.strip()

        # Fenced code block
        if stripped.startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                p = doc.add_paragraph(lines[i].rstrip("\n"))
                for run in p.runs:
                    run.font.name = CODE_FONT
                i += 1
            i += 1
            continue

        # Markdown table syntax → reject. Callers should use word_add_table
        # so position / caption / merge are controlled explicitly and a
        # caption written elsewhere doesn't end up duplicated.
        if (
            stripped.startswith("|")
            and stripped.endswith("|")
            and stripped.count("|") >= 2
        ):
            raise ValueError(
                "markdown insert does not handle '| ... |' table syntax — "
                "use word_add_table (or apply_edits op='add_table') for the "
                "table, with explicit position/anchor/caption. Insert text "
                "and tables as separate ops; do not put markdown table rows "
                "inside the text."
            )

        # ATX headings — caps at H6; deeper levels still render as H6.
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = min(len(m.group(1)), 6)
            p = doc.add_heading(m.group(2).strip(), level=level)
            apply_cjk_font_to_para(p, HEADING_FONT)
            i += 1
            continue

        if re.match(r"^[-*+]\s+", stripped):
            content = re.sub(r"^[-*+]\s+", "", stripped)
            try:
                p = doc.add_paragraph(style="List Bullet")
            except KeyError:
                p = doc.add_paragraph()
            add_inline_runs(p, content, BODY_FONT)
            apply_cjk_font_to_para(p, BODY_FONT)
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            content = re.sub(r"^\d+\.\s+", "", stripped)
            try:
                p = doc.add_paragraph(style="List Number")
            except KeyError:
                p = doc.add_paragraph()
            add_inline_runs(p, content, BODY_FONT)
            apply_cjk_font_to_para(p, BODY_FONT)
            i += 1
            continue

        if stripped == "":
            # Blank line = block separator, not content. Emitting an empty
            # paragraph here leaves a stray blank line between every paragraph
            # (the "每段中间空一格" bug) and diverges from how `create` renders.
            i += 1
            continue
        p = doc.add_paragraph()
        add_inline_runs(p, line, BODY_FONT)
        apply_cjk_font_to_para(p, BODY_FONT)
        i += 1

    all_blocks = [c for c in body if c.tag in (p_tag, tbl_tag)]
    return all_blocks[before_blocks:]
