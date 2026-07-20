"""Markdown → DOCX bytes engine.

The core conversion used to live inside ``mcp_servers/report_export_mcp/impl.py``
but logically belongs in the engine so the word-editing skill (and any other
consumer) can reach it without depending on a sibling MCP server's internals.

Public API:
    markdown_to_docx_bytes(markdown, title) -> bytes
        Convert a Markdown source into a self-contained .docx blob with
        Chinese-document defaults (方正 fonts, 1.5 line spacing, 2-char
        first-line indent on body paragraphs). Uses pandoc when available
        for fidelity; falls back to a pure python-docx renderer that
        handles headings, paragraphs, lists, fenced code blocks, basic
        inline ``**bold**``/``*italic*``/``\\`code\\```/``[link](url)``,
        and pipe tables.
    harmonize_to_chinese_style(docx_bytes) -> bytes
        Re-apply the same 方正-font / 公文 paragraph formatting / black
        headings / centered tables to an existing .docx. Used to make the
        ``create --content`` structural (.NET) output render identically to
        ``create --markdown`` output.

The legacy underscore-prefixed name ``_markdown_to_docx_bytes`` is kept as
an alias for backward compat.
"""
from __future__ import annotations

import contextlib
import io
import logging
import os
import pathlib
import re
import shutil
import subprocess
import tempfile
from typing import Any, List

logger = logging.getLogger(__name__)

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

_BODY_FONT = "方正仿宋简体"
_HEADING_FONT = "方正小标宋简体"
_CODE_FONT = "Courier New"

# Optional pandoc reference template — same lookup as the original
# report_export_mcp version, but here we don't bundle one alongside this
# module, so the lookup almost always falls through (which the code below
# handles gracefully with ``os.path.exists`` checks).
_REFERENCE_DOCX = os.path.join(os.path.dirname(__file__), "reference.docx")

# Table styling constants
_TABLE_HEADER_COLOR = "366092"
_TABLE_ALT_ROW_COLOR = "DCE6F1"
_TABLE_BORDER_COLOR = "B8CCE4"

# Pre-compiled regex for inline markdown parsing
_INLINE_RE = re.compile(
    r"(\*\*(.+?)\*\*)"        # group 1,2: bold
    r"|(\*(.+?)\*)"           # group 3,4: italic
    r"|(`(.+?)`)"             # group 5,6: inline code
    r"|(\[(.+?)\]\((.+?)\))"  # group 7,8,9: link [text](url)
)


# ── DOCX helpers ────────────────────────────────────────────────────────────

def _set_document_default_fonts(doc, ea_font: str) -> None:
    """Write docDefaults into styles.xml so every paragraph inherits the CJK font."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    styles_el = doc.styles.element
    docDefaults = styles_el.find(qn("w:docDefaults"))
    if docDefaults is None:
        docDefaults = OxmlElement("w:docDefaults")
        styles_el.insert(0, docDefaults)

    rPrDefault = docDefaults.find(qn("w:rPrDefault"))
    if rPrDefault is None:
        rPrDefault = OxmlElement("w:rPrDefault")
        docDefaults.append(rPrDefault)

    rPr = rPrDefault.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        rPrDefault.append(rPr)

    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)

    rFonts.set(qn("w:eastAsia"), ea_font)
    rFonts.set(qn("w:cs"), ea_font)


def _is_body_paragraph(para) -> bool:
    """True iff ``para`` should get body-text defaults (indent / spacing / justify)."""
    from docx.oxml.ns import qn

    style_name = (para.style.name or "") if para.style else ""
    if style_name.startswith("Heading") or style_name in ("Title", "Subtitle"):
        return False
    if style_name.startswith("List") or "Bullet" in style_name or "Number" in style_name:
        return False
    if "Code" in style_name or "Verbatim" in style_name:
        return False
    if para.runs and all((r.font.name or "") == _CODE_FONT for r in para.runs if r.text):
        return False
    pPr = para._element.find(qn("w:pPr"))
    if pPr is not None and pPr.find(qn("w:numPr")) is not None:
        return False
    return True


def _apply_body_paragraph_defaults(para) -> None:
    """Apply the default Chinese-report body-paragraph formatting:
    first-line indent 2 chars + 1.5 line spacing + justified alignment.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = 1.5

    pPr = para._element.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(qn("w:firstLineChars"), "200")
    if ind.get(qn("w:firstLine")) is not None:
        del ind.attrib[qn("w:firstLine")]


def _apply_cjk_font_to_para(para, font_name: str) -> None:
    """Apply CJK font to all runs + paragraph mark in a paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _set_rFonts(rPr_el: Any, name: str) -> None:
        rFonts = rPr_el.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr_el.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), name)
        rFonts.set(qn("w:ascii"), name)
        rFonts.set(qn("w:hAnsi"), name)

    for run in para.runs:
        run.font.name = font_name
        _set_rFonts(run._element.get_or_add_rPr(), font_name)

    pPr = para._element.get_or_add_pPr()
    pRpr = pPr.find(qn("w:rPr"))
    if pRpr is None:
        pRpr = OxmlElement("w:rPr")
        pPr.append(pRpr)
    _set_rFonts(pRpr, font_name)


def _setup_heading_styles(doc) -> None:
    """Set CJK font and force black color on heading style definitions.

    Covers Title / Subtitle / Heading 1-6 so no title or sub-heading inherits
    the blue theme color baked into python-docx's (and pandoc's) default
    template. Styles are matched by display name while iterating ``doc.styles``
    rather than ``doc.styles[name]`` — pandoc emits heading style names that
    don't round-trip through python-docx's built-in name lookup, which used to
    silently skip every heading and leave them blue.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import RGBColor

    targets = {"Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3",
               "Heading 4", "Heading 5", "Heading 6"}
    for style in doc.styles:
        if (style.name or "") not in targets:
            continue
        try:
            style.font.name = _HEADING_FONT
            rPr = style.element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), _HEADING_FONT)
            style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            color_el = rPr.find(qn("w:color"))
            if color_el is not None:
                color_el.set(qn("w:val"), "000000")
                for attr in (qn("w:themeColor"), qn("w:themeShade"), qn("w:themeTint")):
                    if attr in color_el.attrib:
                        del color_el.attrib[attr]
        except (KeyError, ValueError):
            continue


def _pandoc_available() -> bool:
    return shutil.which("pandoc") is not None


def _pandoc_convert(markdown: str, title: str) -> bytes:
    """Convert Markdown to DOCX bytes via pandoc subprocess."""
    first_line = markdown.lstrip().split("\n")[0] if markdown.strip() else ""
    already_has_title = first_line.startswith("#")
    full_md = markdown if (not title or already_has_title) else f"# {title}\n\n{markdown}"

    with tempfile.NamedTemporaryFile(
        suffix=".md", mode="w", encoding="utf-8", delete=False
    ) as f_in:
        f_in.write(full_md)
        in_path = f_in.name

    out_path = str(pathlib.Path(in_path).with_suffix(".docx"))
    try:
        cmd = [
            "pandoc", in_path,
            "-f", "gfm",
            "-t", "docx",
            "--wrap=none",
            "-o", out_path,
        ]
        if os.path.exists(_REFERENCE_DOCX):
            cmd += ["--reference-doc", _REFERENCE_DOCX]
        subprocess.run(cmd, check=True, capture_output=True, timeout=30)
        with open(out_path, "rb") as f:
            return f.read()
    finally:
        for p in (in_path, out_path):
            with contextlib.suppress(OSError):
                os.unlink(p)


def _post_process_cjk_fonts(docx_bytes: bytes) -> bytes:
    """Open pandoc-generated docx and ensure CJK fonts and black headings are applied."""
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    _set_document_default_fonts(doc, _BODY_FONT)
    _setup_heading_styles(doc)

    for para in doc.paragraphs:
        is_body = _is_body_paragraph(para)
        font = _HEADING_FONT if para.style.name.startswith("Heading") else _BODY_FONT
        _apply_cjk_font_to_para(para, font)
        if is_body:
            _apply_body_paragraph_defaults(para)

    _style_tables(doc.tables)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _style_tables(tables) -> None:
    """Add header background, alternating row colors, borders, and center each table."""
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import RGBColor

    def _set_cell_shading(cell, color: str) -> None:
        tc_pr = cell._element.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), color)
        tc_pr.append(shading)

    def _set_cell_borders(cell) -> None:
        tc_pr = cell._element.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), _TABLE_BORDER_COLOR)
            borders.append(el)
        tc_pr.append(borders)

    for table in tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                _set_cell_borders(cell)
                if row_idx == 0:
                    _set_cell_shading(cell, _TABLE_HEADER_COLOR)
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                elif row_idx % 2 == 0:
                    _set_cell_shading(cell, _TABLE_ALT_ROW_COLOR)

                for para in cell.paragraphs:
                    _apply_cjk_font_to_para(para, _BODY_FONT)


def _add_inline_runs(paragraph, text: str, font_name: str) -> None:
    """Parse inline markdown (bold, italic, inline code, links) and add runs."""
    from docx.shared import RGBColor

    last_end = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > last_end:
            run = paragraph.add_run(text[last_end:m.start()])
            run.font.name = font_name

        if m.group(2):  # bold
            run = paragraph.add_run(m.group(2))
            run.bold = True
            run.font.name = font_name
        elif m.group(4):  # italic
            run = paragraph.add_run(m.group(4))
            run.italic = True
            run.font.name = font_name
        elif m.group(6):  # inline code
            run = paragraph.add_run(m.group(6))
            run.font.name = _CODE_FONT
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif m.group(8):  # link
            run = paragraph.add_run(m.group(8))
            run.font.name = font_name
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.underline = True

        last_end = m.end()

    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.name = font_name


def _is_separator_row(cells: List[str]) -> bool:
    """Return True iff every non-empty cell looks like a ``|---|:---:|---|`` separator."""
    non_empty = [c.strip() for c in cells if c.strip()]
    if not non_empty:
        return False
    return all(bool(re.match(r"^:?-+:?$", c)) for c in non_empty)


def _table_lines_to_rows(raw_lines: List[str]) -> List[List[str]]:
    """Parse raw markdown table lines into rows of cell strings, skipping separators."""
    rows: List[List[str]] = []
    for line in raw_lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        if _is_separator_row(cells):
            continue
        rows.append(cells)
    return rows


def _add_table_from_lines(doc, table_lines: List[str]) -> None:
    """Parse markdown table lines and add a Word table to the document."""
    rows = _table_lines_to_rows(table_lines)
    if not rows:
        return

    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)

    for row_idx, row_cells in enumerate(rows):
        for col_idx in range(max_cols):
            cell = table.cell(row_idx, col_idx)
            value = row_cells[col_idx] if col_idx < len(row_cells) else ""
            cell.text = ""
            para = cell.paragraphs[0]
            _add_inline_runs(para, value, _BODY_FONT)
            _apply_cjk_font_to_para(para, _BODY_FONT)


def _fallback_markdown_to_docx(markdown: str, title: str) -> bytes:
    """Pure python-docx fallback when pandoc is unavailable."""
    from docx import Document

    doc = Document()
    _set_document_default_fonts(doc, _BODY_FONT)
    _setup_heading_styles(doc)

    lines = (markdown or "").splitlines()
    first_non_empty = next((l.strip() for l in lines if l.strip()), "")
    if not first_non_empty.startswith("#"):
        title_para = doc.add_heading(title or "报告", level=0)
        _apply_cjk_font_to_para(title_para, _HEADING_FONT)
    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")
        stripped = line.strip()

        if stripped.startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                p = doc.add_paragraph(lines[i].rstrip("\n"))
                for run in p.runs:
                    run.font.name = _CODE_FONT
                i += 1
            i += 1
            continue

        if (
            stripped.startswith("|")
            and stripped.endswith("|")
            and stripped.count("|") >= 2
        ):
            table_lines = []
            while i < len(lines):
                s = lines[i].strip()
                if s.startswith("|") and s.endswith("|") and s.count("|") >= 2:
                    table_lines.append(s)
                    i += 1
                else:
                    break
            _add_table_from_lines(doc, table_lines)
            continue

        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:].strip(), level=3)
            _apply_cjk_font_to_para(p, _HEADING_FONT)
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:].strip(), level=2)
            _apply_cjk_font_to_para(p, _HEADING_FONT)
        elif stripped.startswith("# "):
            p = doc.add_heading(stripped[2:].strip(), level=1)
            _apply_cjk_font_to_para(p, _HEADING_FONT)
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, stripped[2:].strip(), _BODY_FONT)
            _apply_cjk_font_to_para(p, _BODY_FONT)
        elif re.match(r"^\d+\. ", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, re.sub(r"^\d+\. ", "", stripped), _BODY_FONT)
            _apply_cjk_font_to_para(p, _BODY_FONT)
        else:
            p = doc.add_paragraph()
            _add_inline_runs(p, line, _BODY_FONT)
            _apply_cjk_font_to_para(p, _BODY_FONT)
            if stripped:
                _apply_body_paragraph_defaults(p)

        i += 1

    _style_tables(doc.tables)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _has_hanging_indent(para) -> bool:
    """True iff the paragraph carries a hanging indent.

    The structural (.NET) engine renders list items as plain paragraphs with a
    ``<w:ind w:hanging=...>``. Such paragraphs must NOT get the body-text
    first-line indent, which would conflict with the hanging indent.
    """
    from docx.oxml.ns import qn

    pPr = para._element.find(qn("w:pPr"))
    if pPr is None:
        return False
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        return False
    return ind.get(qn("w:hanging")) is not None or ind.get(qn("w:hangingChars")) is not None


def harmonize_to_chinese_style(docx_bytes: bytes) -> bytes:
    """Re-apply the canonical Chinese-document styling to an existing .docx.

    Post-processes documents built by the .NET structural engine
    (``create --content``) so they render with the *same* fonts and formatting
    as the markdown engine output (``create --markdown``): 方正 fonts, black
    headings, justified body text with 1.5 line spacing and a 2-character
    first-line indent, and centered tables.
    """
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    _set_document_default_fonts(doc, _BODY_FONT)
    _setup_heading_styles(doc)

    for para in doc.paragraphs:
        style_name = (para.style.name or "") if para.style else ""
        is_heading = style_name.startswith("Heading") or style_name in ("Title", "Subtitle")
        _apply_cjk_font_to_para(para, _HEADING_FONT if is_heading else _BODY_FONT)
        if _is_body_paragraph(para) and not _has_hanging_indent(para):
            _apply_body_paragraph_defaults(para)

    # Header / footer paragraphs share the body font for visual consistency.
    for section in doc.sections:
        for hf in (section.header, section.footer):
            if hf.is_linked_to_previous:
                continue
            for para in hf.paragraphs:
                _apply_cjk_font_to_para(para, _BODY_FONT)

    _style_tables(doc.tables)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def markdown_to_docx_bytes(markdown: str, title: str) -> bytes:
    """Convert Markdown to DOCX bytes. Uses pandoc if available, otherwise pure python-docx."""
    if _pandoc_available():
        logger.info("Using pandoc for markdown-to-docx conversion")
        raw = _pandoc_convert(markdown, title)
        return _post_process_cjk_fonts(raw)
    logger.info("Pandoc not available, using fallback python-docx converter")
    return _fallback_markdown_to_docx(markdown, title)


# Backward-compat alias: the old call site in mcp_servers/report_export_mcp/impl.py
# imported the underscore-prefixed name.
_markdown_to_docx_bytes = markdown_to_docx_bytes
