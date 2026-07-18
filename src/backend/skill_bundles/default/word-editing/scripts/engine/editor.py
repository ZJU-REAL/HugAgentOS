"""Word document editing primitives.

Design notes
============

**Run-preservation policy** (the #1 source of "edit succeeded but formatting
broke" bugs in the prior implementation):

    python-docx represents a paragraph's visible text as a sequence of
    ``Run`` objects, each with its own formatting. Word frequently splits
    runs at formatting boundaries, language toggles, spell-check anchors,
    etc., so a literal ``find`` string can be present in the visible
    paragraph text yet not in any single run's ``.text``.

    This module's ``_replace_in_paragraph`` uses a two-tier strategy:

    1. **per-run fast path** — try literal replace within each run's text
       independently. When the match falls entirely inside a single run,
       only that run is mutated; every other run keeps its formatting
       intact. This covers the vast majority of "fix one number / swap a
       word" edits.

    2. **paragraph-level fallback** — only invoked when the per-run path
       finds nothing AND a match exists in the concatenated paragraph
       text (i.e. the match crosses run boundaries) or regex is enabled.
       The matched span's secondary inline formatting is collapsed into
       the first run; this is the documented trade-off for cross-run
       edits and matches the GongRzhe/Office-Word-MCP-Server behavior.

    After every mutation the affected paragraph runs through
    ``_merge_adjacent_runs`` which fuses neighboring runs that share the
    same ``rPr`` — keeps run count from inflating across many edits.

**Lenient matching** (``lenient=True``):

    Normalizes both haystack and needle through NFKC, strips zero-width /
    soft-hyphen / BOM chars, and treats any whitespace span (incl. NBSP,
    full-width space) as equivalent. Forces the paragraph-rebuild fallback
    because matched span boundaries don't survive normalization. Use only
    when literal find genuinely fails — it loses inline formatting in the
    matched span.

**Atomic editing** (``apply_edits``):

    Opens the document once, runs an ordered list of operations, saves
    once. Per-op result is reported back. This is the path the LLM should
    prefer for any multi-edit workflow — it avoids N×open/save cycles AND
    eliminates the paragraph-index drift problem (downstream ops see the
    in-memory state mutated by upstream ops, no need to re-read outline).
"""
from __future__ import annotations

import re
import unicodedata
from typing import Any, Callable, Iterable, Optional

from ._handle import input_path, output_path

# ── normalization ───────────────────────────────────────────────────────────

# Zero-width / formatting code points that copy-paste / OCR / Word toggles
# inject into otherwise-equal-looking strings. Stripping them gives lenient
# matching the best chance.
_ZERO_WIDTH = "".join([
    "​",  # ZWSP
    "‌",  # ZWNJ
    "‍",  # ZWJ
    "⁠",  # word joiner
    "﻿",  # BOM / ZWNBSP
    "­",  # soft hyphen
])
_ZERO_WIDTH_RE = re.compile(f"[{_ZERO_WIDTH}]")
_WHITESPACE_RE = re.compile(r"\s+")


def _normalize_lenient(text: str) -> str:
    """NFKC + strip zero-width + collapse whitespace.

    Used by ``lenient=True`` matching. Result is suitable for substring /
    regex search but NOT for direct re-insertion (it loses original
    spacing). Callers must still rebuild text from the original runs.
    """
    if not text:
        return ""
    out = unicodedata.normalize("NFKC", text)
    out = _ZERO_WIDTH_RE.sub("", out)
    out = _WHITESPACE_RE.sub(" ", out).strip()
    return out


# ── traversal ───────────────────────────────────────────────────────────────


def _iter_all_paragraphs(doc) -> Iterable[Any]:
    """Yield every paragraph in body, tables, headers, footers."""
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in doc.sections:
        for hdr_or_ftr in (section.header, section.footer):
            if hdr_or_ftr is None:
                continue
            yield from hdr_or_ftr.paragraphs
            for table in hdr_or_ftr.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from cell.paragraphs


def _iter_body_paragraphs(doc) -> list[Any]:
    """Body-only paragraphs (matches reader.get_outline indexing)."""
    return list(doc.paragraphs)


# ── run merging (housekeeping after edits) ──────────────────────────────────


def _rpr_xml(run) -> str:
    """Stringify a run's rPr (or empty string if none) for equality compare.

    python-docx exposes ``run.element`` as an lxml node; we serialize the
    rPr child to a stable string. Cheap enough — paragraphs rarely have
    > a few dozen runs.
    """
    el = run._element
    rpr = el.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr")
    if rpr is None:
        return ""
    from lxml import etree
    return etree.tostring(rpr).decode("utf-8")


def _merge_adjacent_runs(para) -> int:
    """Merge neighboring runs that share identical rPr.

    Returns the number of runs removed. python-docx represents text as
    ``<w:r><w:t>...</w:t></w:r>`` so we mutate the parent paragraph's
    XML directly: append cur's text to prev's <w:t>, drop cur.
    """
    runs = list(para.runs)
    if len(runs) < 2:
        return 0
    removed = 0
    i = 1
    while i < len(runs):
        prev = runs[i - 1]
        cur = runs[i]
        # Only merge plain text runs — skip if either has fields, drawings,
        # tabs, breaks, or other non-text children we'd lose by concatenating.
        if not _is_plain_text_run(prev) or not _is_plain_text_run(cur):
            i += 1
            continue
        if _rpr_xml(prev) != _rpr_xml(cur):
            i += 1
            continue
        prev.text = (prev.text or "") + (cur.text or "")
        cur._element.getparent().remove(cur._element)
        runs.pop(i)
        removed += 1
    return removed


def _is_plain_text_run(run) -> bool:
    """True iff the run only contains rPr + w:t children (safe to merge)."""
    w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    for child in run._element:
        tag = child.tag
        if tag in (w_ns + "rPr", w_ns + "t"):
            continue
        return False
    return True


# ── core replace ────────────────────────────────────────────────────────────


def _replace_in_paragraph(
    para,
    pattern: str | re.Pattern,
    replacement: str,
    *,
    max_replacements: int,
    lenient: bool = False,
) -> int:
    """Replace occurrences inside one paragraph; return number replaced.

    ``max_replacements=-1`` means replace all in this paragraph.

    Strategy:
        1. **Per-run fast path** (literal only): try each run independently.
           If the find string fits inside one run, only that run is mutated.
           Other runs' formatting is preserved 100%.
        2. **Paragraph rebuild fallback**: if no per-run hit but the
           concatenated paragraph text matches (or regex/lenient is on),
           rebuild via collapse-into-first-run. Loses secondary inline
           formatting in the matched span.

    Always runs ``_merge_adjacent_runs`` after a successful mutation so
    repeated edits don't bloat run counts.
    """
    if not para.runs:
        return 0
    if max_replacements == 0:
        return 0

    full_text = para.text

    # ── per-run fast path (literal only, non-lenient) ──
    if isinstance(pattern, str) and not lenient:
        if pattern not in full_text:
            return 0
        n_total = 0
        remaining = max_replacements
        for run in para.runs:
            if remaining == 0:
                break
            run_text = run.text or ""
            if pattern not in run_text:
                continue
            occurrences = run_text.count(pattern)
            if remaining < 0:
                run.text = run_text.replace(pattern, replacement)
                n_total += occurrences
            else:
                take = min(occurrences, remaining)
                run.text = run_text.replace(pattern, replacement, take)
                n_total += take
                remaining -= take
        if n_total > 0:
            _merge_adjacent_runs(para)
            return n_total
        # No per-run hit but the literal IS in para.text — match crosses
        # run boundaries. Fall through to paragraph rebuild.

    # ── paragraph-level rebuild fallback ──
    if lenient and isinstance(pattern, str):
        # Build a normalized haystack; find normalized needle in it. If
        # found, fall back to literal replacement in the original paragraph
        # text using the normalized form as a search proxy. We can't do
        # surgical run editing under normalization, so this collapses runs.
        norm_full = _normalize_lenient(full_text)
        norm_needle = _normalize_lenient(pattern)
        if not norm_needle or norm_needle not in norm_full:
            return 0
        # Build a regex from the normalized needle that tolerates the same
        # ambiguities in the original (raw) paragraph text:
        # - any run of whitespace == \s+
        # - zero-width chars allowed between any two non-space chars
        regex = _build_lenient_regex(pattern)
        if max_replacements < 0:
            new_text, n = regex.subn(replacement, full_text)
        else:
            new_text, n = regex.subn(replacement, full_text, count=max_replacements)
        if n == 0 or new_text == full_text:
            return 0
    elif isinstance(pattern, re.Pattern):
        if not pattern.search(full_text):
            return 0
        if max_replacements < 0:
            new_text, n = pattern.subn(replacement, full_text)
        else:
            new_text, n = pattern.subn(replacement, full_text, count=max_replacements)
        if n == 0 or new_text == full_text:
            return 0
    else:
        # Literal cross-run match (per-run path didn't catch it).
        if pattern not in full_text:
            return 0
        if max_replacements < 0:
            new_text = full_text.replace(pattern, replacement)
            n = full_text.count(pattern)
        else:
            n = min(full_text.count(pattern), max_replacements)
            new_text = full_text.replace(pattern, replacement, n)
        if n == 0:
            return 0

    # Collapse: keep first run's rPr, drop secondary runs' content.
    first_run = para.runs[0]
    first_run.text = new_text
    for r in para.runs[1:]:
        r.text = ""
    _merge_adjacent_runs(para)
    return n


def _build_lenient_regex(needle: str) -> re.Pattern:
    """Build a regex from ``needle`` that matches the lenient-equivalent text.

    - Whitespace runs in ``needle`` → ``[\\s + zero-width]+``: matches NBSP,
      full-width space, tab, AND tolerates zero-width chars interleaved.
    - Other characters: escaped literally, with optional zero-width chars
      allowed between them so e.g. "foo" matches "f<ZWSP>oo".
    - The pattern is anchored to whole-grapheme matches; inline char-class
      gaps stay tiny so length grows linearly with needle length.
    """
    out: list[str] = []
    zw_gap = f"[{_ZERO_WIDTH}]*"
    ws_class = f"[\\s{_ZERO_WIDTH}]+"
    parts = re.split(r"(\s+)", needle)
    for part in parts:
        if not part:
            continue
        if part.isspace():
            out.append(ws_class)
            continue
        # Escape each char and join with zero-width tolerance.
        chars = [re.escape(c) for c in part]
        out.append(zw_gap.join(chars))
    return re.compile("".join(out))


# ── primitives operating on an open Document ────────────────────────────────
#
# The `_op_*` functions all take an open ``doc`` and mutate it in place.
# They are the building blocks for both the single-shot public API
# (search_and_replace, insert_text, …) and the batch ``apply_edits`` path.


def _op_search_and_replace(
    doc,
    *,
    find: str,
    replace: str,
    scope: str = "all",
    regex: bool = False,
    lenient: bool = False,
) -> dict[str, Any]:
    if scope not in ("first", "all"):
        raise ValueError(f"scope must be 'first' or 'all', got {scope!r}")
    if not find:
        raise ValueError("'find' must be a non-empty string")

    pattern: str | re.Pattern
    if regex:
        try:
            pattern = re.compile(find)
        except re.error as e:
            raise ValueError(f"invalid regex: {e}") from e
    else:
        pattern = find

    quota = 1 if scope == "first" else -1
    done = 0
    for para in _iter_all_paragraphs(doc):
        if quota == 0:
            break
        remaining = quota if quota > 0 else -1
        n = _replace_in_paragraph(
            para, pattern, replace,
            max_replacements=remaining,
            lenient=lenient and not regex,
        )
        done += n
        if quota > 0:
            quota -= n

    out: dict[str, Any] = {"replacements": done, "scope": scope, "regex": regex, "lenient": lenient}
    if done == 0:
        # Cross-paragraph find is the #1 reason chained edits fail. The replace
        # algorithm operates per-paragraph (matching `<w:p>` boundaries),
        # so a `find` containing a literal LF can never match — the LF lives in
        # the get_document_text JOIN, not in any single paragraph element.
        if "\n" in find:
            out["warning"] = (
                "no matches; 'find' contains a newline (\\n) but search is "
                "PARAGRAPH-SCOPED — `find` is matched within each <w:p> "
                "individually and cannot span paragraphs. To rewrite a "
                "multi-paragraph block, use `word_apply_edits` with "
                "`{op: 'replace_section', heading_anchor: '...'}` or chain "
                "`{op: 'replace_paragraph', anchor: '...', new_text: '...'}` "
                "ops. For an entire chapter rewrite prefer `replace_section`."
            )
        else:
            out["warning"] = (
                f"no matches for find={find!r}; the search string may have "
                "hidden whitespace, zero-width chars, full-width punctuation, "
                "or be split across run boundaries. Try lenient=True, or read "
                "the document text first to confirm the exact substring."
            )
    return out


def _op_replace_many(
    doc,
    *,
    replacements: list[dict[str, Any]],
) -> dict[str, Any]:
    if not isinstance(replacements, list) or not replacements:
        raise ValueError("'replacements' must be a non-empty list")
    per_op: list[dict[str, Any]] = []
    total = 0
    for i, item in enumerate(replacements):
        if not isinstance(item, dict):
            raise ValueError(f"replacements[{i}] must be a dict")
        try:
            sub = _op_search_and_replace(
                doc,
                find=item.get("find") or "",
                replace=item.get("replace") or "",
                scope=item.get("scope", "all"),
                regex=bool(item.get("regex", False)),
                lenient=bool(item.get("lenient", False)),
            )
        except ValueError as e:
            sub = {"ok": False, "error": str(e), "find": item.get("find")}
        else:
            sub["ok"] = True
            sub["find"] = item.get("find")
            total += sub.get("replacements", 0)
        per_op.append(sub)
    return {"replacements": total, "per_op": per_op}


def _op_fill_placeholders(
    doc,
    *,
    mapping: dict[str, str],
    pattern: str = r"\{\{(\w+)\}\}",
) -> dict[str, Any]:
    if not isinstance(mapping, dict) or not mapping:
        raise ValueError("'mapping' must be a non-empty dict[str, str]")
    try:
        regex = re.compile(pattern)
    except re.error as e:
        raise ValueError(f"invalid pattern: {e}") from e

    filled: dict[str, int] = {k: 0 for k in mapping}
    unfilled: dict[str, int] = {}

    for para in _iter_all_paragraphs(doc):
        text = para.text
        if not text:
            continue
        # Find all {{name}} hits in this paragraph; replace those that map.
        # Each hit goes through _replace_in_paragraph so single-run hits
        # preserve formatting.
        for m in list(regex.finditer(text)):
            full = m.group(0)
            key = m.group(1) if regex.groups >= 1 else full
            if key in mapping:
                n = _replace_in_paragraph(
                    para, full, mapping[key], max_replacements=-1
                )
                if n > 0:
                    filled[key] = filled.get(key, 0) + n
                else:
                    # match disappeared between finditer and replace (e.g.
                    # a previous iteration mutated the paragraph text); skip.
                    pass
            else:
                unfilled[key] = unfilled.get(key, 0) + 1

    return {
        "filled": {k: v for k, v in filled.items() if v > 0},
        "unfilled_keys": sorted(unfilled.keys()),
        "unmapped_placeholders_seen": sum(unfilled.values()),
    }


def _op_list_placeholders(
    doc,
    *,
    pattern: str = r"\{\{(\w+)\}\}",
) -> dict[str, Any]:
    try:
        regex = re.compile(pattern)
    except re.error as e:
        raise ValueError(f"invalid pattern: {e}") from e

    counts: dict[str, int] = {}
    for para in _iter_all_paragraphs(doc):
        text = para.text
        if not text:
            continue
        for m in regex.finditer(text):
            key = m.group(1) if regex.groups >= 1 else m.group(0)
            counts[key] = counts.get(key, 0) + 1
    return {
        "placeholders": sorted(counts.keys()),
        "counts": counts,
        "total_occurrences": sum(counts.values()),
    }


# ── Markdown helpers re-exported from word.markdown ─────────────────────────
# Implementation lives in ``the engine/word/markdown.py``; we re-export the
# names editor.py historically used so callers (and tests reaching into
# private symbols via the legacy names) keep working.

from .markdown import (  # noqa: E402
    INLINE_RE as _MD_INLINE_RE,
    add_inline_runs as _md_add_inline_runs,
    emit_blocks as _emit_markdown_blocks,
    looks_like_markdown as _looks_like_markdown,
)


# ── shared content-format resolution for insert / rewrite ops ────────────────
# ``insert`` / ``replace_paragraph`` / ``replace_section`` all take a free-text
# content blob. The LLM routinely authors that blob as Markdown (``###`` /
# ``**bold**`` / ``- `` lists). Rendering it literally is the #1 "the docx
# shows '### 标题' verbatim" complaint, so these ops share one format gate.

_CONTENT_FORMATS = ("text", "markdown", "auto")

# Caller-facing hint emitted when a content op runs in literal ``text`` mode on
# input that clearly carries Markdown markup (only reachable when the caller
# explicitly forced ``format="text"``; ``auto`` would have rendered it).
_MARKDOWN_LITERAL_WARNING = (
    "input contains markdown signals (e.g. '###', '**', '- ', '|...|') but "
    "format='text' inserts them as literal characters — the docx will show "
    "'### 标题' verbatim instead of rendering. Drop format (defaults to 'auto') "
    "or pass format='markdown' to render headings/bold/lists."
)


def _resolve_content_format(text: str, fmt: str) -> str:
    """Normalize a content-op ``format`` value to ``'text'`` or ``'markdown'``.

    ``"auto"`` (the default for the content ops) renders as Markdown only when
    the content actually carries Markdown signals; plain prose stays literal so
    existing plain-text edits are byte-for-byte unchanged. ``"markdown"`` /
    ``"text"`` force the respective path.
    """
    if fmt not in _CONTENT_FORMATS:
        raise ValueError(
            f"format must be one of {_CONTENT_FORMATS}, got {fmt!r}"
        )
    if fmt == "auto":
        return "markdown" if _looks_like_markdown(text or "") else "text"
    return fmt


def _emit_markdown_blocks_safe(doc, markdown: str) -> list[Any]:
    """``_emit_markdown_blocks`` with rollback of any partially-appended blocks.

    ``_emit_markdown_blocks`` appends blocks to the body as it parses and may
    raise partway through (e.g. on ``| … |`` table syntax). The rewrite ops
    emit BEFORE they delete the old content, so on failure the document must be
    left exactly as it was — otherwise half-rendered orphan paragraphs leak to
    the end of the body. This wrapper removes anything the failed emit appended
    and re-raises the original error untouched.
    """
    from docx.oxml.ns import qn

    body = doc.element.body
    block_tags = (qn("w:p"), qn("w:tbl"))
    before = {id(c) for c in body if c.tag in block_tags}
    try:
        return _emit_markdown_blocks(doc, markdown)
    except Exception:
        for el in [c for c in body if c.tag in block_tags]:
            if id(el) not in before:
                parent = el.getparent()
                if parent is not None:
                    parent.remove(el)
        raise


def _relocate_after(anchor_el, new_elements: list[Any]) -> None:
    """Move ``new_elements`` (document order) to immediately after ``anchor_el``."""
    cur = anchor_el
    for el in new_elements:
        cur.addnext(el)
        cur = el


def _relocate_before(target_el, new_elements: list[Any]) -> None:
    """Move ``new_elements`` (document order) to immediately before ``target_el``."""
    for el in new_elements:
        target_el.addprevious(el)


def _count_blocks(new_elements: list[Any]) -> tuple[int, int]:
    """Return ``(paragraph_count, table_count)`` for a list of OOXML blocks."""
    from docx.oxml.ns import qn

    p_tag, tbl_tag = qn("w:p"), qn("w:tbl")
    n_p = sum(1 for el in new_elements if el.tag == p_tag)
    n_tbl = sum(1 for el in new_elements if el.tag == tbl_tag)
    return n_p, n_tbl


# ── body-paragraph format inheritance (keep edits looking like the doc) ──────
# `create` bakes Chinese-report body formatting (first-line-indent 2 chars via
# ``firstLineChars=200`` + 1.5 line spacing + justify) DIRECTLY onto each
# paragraph — the Normal style carries none of it. So a fresh
# ``doc.add_paragraph()`` from a rewrite op loses the indent/spacing and looks
# different from the surrounding text. These helpers snapshot the paragraph-
# level format (``w:ind`` / ``w:spacing`` / ``w:jc``) from the content being
# replaced and replay it onto the new paragraphs so the edit matches its
# neighbours. We copy the raw OOXML because ``firstLineChars`` is char-based and
# python-docx's ``first_line_indent`` (EMU-only) can't even read it.

_BODY_FMT_TAGS = ("w:spacing", "w:ind", "w:jc")


def _capture_body_format(para) -> list[Any]:
    """Snapshot ``w:ind`` / ``w:spacing`` / ``w:jc`` from ``para`` as clones."""
    from copy import deepcopy
    from docx.oxml.ns import qn

    pPr = para._element.find(qn("w:pPr"))
    if pPr is None:
        return []
    caps: list[Any] = []
    for tag in _BODY_FMT_TAGS:
        el = pPr.find(qn(tag))
        if el is not None:
            caps.append(deepcopy(el))
    return caps


def _is_body_block(el) -> bool:
    """True iff a ``<w:p>`` should receive body formatting — i.e. it is not a
    heading / title / list item. Reads the raw ``w:pStyle`` id so it serves both
    Paragraph-API and raw-OOXML callers from one predicate."""
    from docx.oxml.ns import qn

    pPr = el.find(qn("w:pPr"))
    if pPr is None:
        return True
    if pPr.find(qn("w:numPr")) is not None:  # a list item
        return False
    pStyle = pPr.find(qn("w:pStyle"))
    val = pStyle.get(qn("w:val")) if pStyle is not None else None
    return not (val and (val.startswith(("Heading", "Title", "Subtitle")) or "List" in val))


def _scan_body_format(doc, *, skip: Optional[set] = None) -> list[Any]:
    """Fallback reference: first body paragraph in the doc that carries a
    ``w:ind`` (i.e. an indented body paragraph). Used when the replaced section
    had no body of its own to copy from."""
    from docx.oxml.ns import qn

    skip = skip or set()
    for p in doc.paragraphs:
        if id(p._element) in skip or not _is_body_block(p._element):
            continue
        pPr = p._element.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:ind")) is not None:
            return _capture_body_format(p)
    return []


def _apply_body_format(el_or_para, caps: list[Any]) -> None:
    """Replay captured ``w:ind`` / ``w:spacing`` / ``w:jc`` onto one paragraph
    (a Paragraph OR a raw ``<w:p>`` element), skipping headings / list items.
    Each application gets its own clone, since an lxml element has one parent."""
    from copy import deepcopy
    from docx.oxml.ns import qn

    if not caps:
        return
    el = getattr(el_or_para, "_element", el_or_para)
    if not _is_body_block(el):
        return
    pPr = el.find(qn("w:pPr"))
    if pPr is None:
        pPr = el.makeelement(qn("w:pPr"), {})
        el.insert(0, pPr)
    for cap in caps:
        existing = pPr.find(cap.tag)
        if existing is not None:
            pPr.remove(existing)
        pPr.append(deepcopy(cap))


def _apply_body_format_to_blocks(new_elements: list[Any], caps: list[Any]) -> None:
    """Replay captured body format onto the ``<w:p>`` body paragraphs among raw
    OOXML ``new_elements`` (skip headings / list items / tables)."""
    from docx.oxml.ns import qn

    if not caps:
        return
    p_tag = qn("w:p")
    for el in new_elements:
        if el.tag == p_tag:
            _apply_body_format(el, caps)


def _split_body_lines(text: str) -> list[str]:
    """Split rewrite-op body text into paragraph lines, treating blank lines as
    mere separators (NOT empty paragraphs). Chinese-report convention separates
    paragraphs with the first-line indent, not blank lines — and a markdown-style
    ``\\n\\n`` between paragraphs must not leave an empty paragraph behind. Falls
    back to ``[""]`` so an all-blank / empty input still yields one paragraph."""
    return [ln for ln in (text or "").split("\n") if ln.strip()] or [""]


def _op_insert_text(
    doc,
    *,
    text: str,
    position: str = "end",
    anchor: Optional[str] = None,
    style: Optional[str] = None,
    style_for_all: bool = False,
    format: str = "auto",
) -> dict[str, Any]:
    """Insert paragraph(s) into a document.

    Multi-line ``text`` is split on ``\\n`` into multiple paragraphs.

    Style application policy (the most common LLM intent is "drop a heading
    plus some body content in one call"):

      * ``style`` (when set) applies to the **first paragraph only** by
        default. Subsequent paragraphs use the document's default style
        (Normal) so body content keeps its body font / formatting.
      * Pass ``style_for_all=True`` to override this and apply ``style`` to
        every paragraph (use case: inserting a list of bulleted items that
        all need the same paragraph style).

    Without this distinction, a single ``style="Heading 3"`` call would turn
    every line into a heading — producing body paragraphs in the wrong font.

    ``format`` parameter:
        - ``"auto"`` (default): render as Markdown when the input carries
          markdown signals (``###``, ``- ``, ``**``, ``|x|``, …), otherwise
          treat it as plain text. This is what the LLM almost always wants —
          markdown drafts render, plain prose stays literal — with no flag to
          remember.
        - ``"markdown"``: always parse ``text`` as Markdown. Emits proper Word
          headings (``#`` → Heading 1 … ``######`` → Heading 6), bullet /
          number lists, fenced code blocks, and inline ``**bold**`` /
          ``*italic*`` / `` `code` `` / ``[link](url)``. ``style`` and
          ``style_for_all`` are IGNORED in markdown mode (block styles come
          from the markdown structure itself).
        - ``"text"``: force plain-text mode. Each ``\\n`` becomes a new
          paragraph; markdown syntax is inserted as literal characters
          (``**bold**`` shows up as ``**bold**`` in the docx). When the input
          looks like markdown a ``warning`` is added to the result.
    """
    from .styles import (
        BODY_FONT,
        apply_cjk_font_to_para,
        font_for_style,
    )

    from .styles import (
        ANCHOR_REQUIRED_POSITIONS as _ANCHOR_REQUIRED,
        INSERT_POSITIONS as _VALID_POSITIONS,
    )
    if position not in _VALID_POSITIONS:
        raise ValueError(
            f"position must be one of {_VALID_POSITIONS}, got {position!r}"
        )
    if position in _ANCHOR_REQUIRED and anchor is None:
        raise ValueError(f"position={position!r} requires 'anchor'")
    if not text:
        raise ValueError("'text' must be non-empty")
    eff_format = _resolve_content_format(text, format)

    # Redirect after_heading→after_section when the new content's first line
    # is a markdown heading at the same/higher level than the anchor heading
    # (i.e. caller is adding a sibling sub-section, not nesting under the
    # anchor). Naive after_heading would wedge the new section between the
    # anchor heading and its own body.
    redirected = False
    if (
        position == "after_heading"
        and eff_format == "markdown"
        and isinstance(anchor, str)
    ):
        first_line = next(
            (ln.strip() for ln in text.splitlines() if ln.strip()), ""
        )
        m = re.match(r"^(#{1,6})\s+", first_line)
        if m:
            new_level = len(m.group(1))
            anchor_level = _heading_level_of_anchor(doc, anchor)
            if anchor_level is not None and new_level <= anchor_level:
                position = "after_section"
                redirected = True

    # ── Markdown branch — parse blocks, then relocate to anchor ──
    if eff_format == "markdown":
        result = _insert_markdown_at_position(
            doc, markdown=text, position=position, anchor=anchor,
        )
        if redirected:
            result["position_redirected_from"] = "after_heading"
            result["redirect_reason"] = (
                "first line is a markdown heading at same/higher level than "
                "the anchor heading — redirected to after_section so the new "
                "sub-section lands at the END of the anchor's section instead "
                "of being wedged between the anchor heading and its own body."
            )
        return result

    body = _iter_body_paragraphs(doc)
    lines = [ln for ln in (text.splitlines() or [text])]
    first_font = font_for_style(style)
    body_font = font_for_style(None)  # Normal-style body font

    def _new_para(line: str, idx: int):
        is_first_or_all = (idx == 0) or style_for_all
        applied_style = style if (style and is_first_or_all) else None
        if applied_style:
            p = doc.add_paragraph(line, style=applied_style)
            apply_cjk_font_to_para(p, first_font)
        else:
            p = doc.add_paragraph(line)
            apply_cjk_font_to_para(p, body_font)
        return p

    if position == "end":
        for i, line in enumerate(lines):
            _new_para(line, i)
    elif position == "start":
        new_paras = [_new_para(line, i) for i, line in enumerate(lines)]
        if body:
            anchor_el = body[0]._element
            for p in new_paras:
                anchor_el.addprevious(p._element)
    elif position == "after_heading":
        target = None
        for p in body:
            sty = p.style.name if p.style else ""
            if sty.startswith(("Title", "Heading")) and (anchor or "") in p.text:
                target = p
                break
        if target is None:
            raise ValueError(f"no heading containing {anchor!r} found")
        # Insert in reverse with addnext so the final order matches input.
        for i, line in enumerate(reversed(lines)):
            new_p = _new_para(line, len(lines) - 1 - i)
            target._element.addnext(new_p._element)
    elif position == "after_paragraph":
        idx = _resolve_paragraph_index(body, anchor)
        target = body[idx]
        # If the anchor paragraph is followed by tables, treat the caption +
        # its table as one unit and advance the insertion point past them.
        # (Otherwise inserting "after the caption" wedges new content
        # BETWEEN caption and table.)
        anchor_el = _advance_past_following_tables(target._element)
        for i, line in enumerate(reversed(lines)):
            new_p = _new_para(line, len(lines) - 1 - i)
            anchor_el.addnext(new_p._element)
    elif position == "before_paragraph":
        idx = _resolve_paragraph_index(body, anchor)
        target = body[idx]
        for i, line in enumerate(lines):
            new_p = _new_para(line, i)
            target._element.addprevious(new_p._element)
    elif position == "after_section":
        # "After the entire section anchored at this heading" — find the
        # heading, then walk to the next paragraph styled Title or Heading
        # at the SAME or HIGHER level (smaller N is higher level), and
        # insert just before it. If there is no such follow-up heading,
        # insert at end of body. Lets the model say "新增一节 3.3" without
        # having to predict the next sibling heading text.
        h_idx, _, content_end = _resolve_section_bounds(body, anchor)
        if content_end < len(body):
            target = body[content_end]
            for i, line in enumerate(lines):
                new_p = _new_para(line, i)
                target._element.addprevious(new_p._element)
        else:
            # No follow-up heading → append at end of body (before sectPr).
            from docx.oxml.ns import qn as _qn
            section_pr = doc.element.body.find(_qn("w:sectPr"))
            for i, line in enumerate(lines):
                new_p = _new_para(line, i)
                if section_pr is not None:
                    section_pr.addprevious(new_p._element)

    result: dict[str, Any] = {
        "inserted_paragraphs": len(lines),
        "position": position,
        "anchor": anchor,
        "style_first_only": bool(style and not style_for_all),
        "style_for_all": bool(style and style_for_all),
        "format": "text",
    }
    # Reachable only when the caller forced format="text" on markdown-looking
    # input (auto would have taken the markdown branch). Warn so the literal
    # rendering isn't a silent surprise.
    if _looks_like_markdown(text):
        result["warning"] = _MARKDOWN_LITERAL_WARNING
    return result


def _insert_markdown_at_position(
    doc, *, markdown: str, position: str, anchor: Optional[str],
) -> dict[str, Any]:
    """Build markdown blocks (using ``_emit_markdown_blocks``) and relocate the
    resulting OOXML elements to the requested position.

    Strategy: ``_emit_markdown_blocks`` appends new ``<w:p>`` / ``<w:tbl>``
    elements to the body in document order and returns them. We then move
    that group to the target anchor via lxml ``addnext`` / ``addprevious``,
    preserving order.
    """
    from docx.oxml.ns import qn

    # 1. resolve target paragraphs BEFORE emit (snapshot of pre-existing body).
    # ``_iter_body_paragraphs`` returns Paragraph objects; we capture them now
    # so the anchor search runs against the doc as the caller knows it,
    # uncontaminated by the markdown blocks we're about to append.
    original_paragraphs = list(_iter_body_paragraphs(doc))

    # 2. emit (appends new <w:p> / <w:tbl> to body end)
    new_elements = _emit_markdown_blocks(doc, markdown)

    target_el = None
    if position == "end":
        # Already at end, but if there's a sectPr we must keep it last.
        body = doc.element.body
        sect = body.find(qn("w:sectPr"))
        if sect is not None:
            for el in new_elements:
                sect.addprevious(el)
        # else: nothing to do (already appended in document order)
    elif position == "start":
        if original_paragraphs:
            target_el = original_paragraphs[0]._element
            # Insert in REVERSE so addprevious produces forward order.
            for el in reversed(new_elements):
                target_el.addprevious(el)
    elif position == "after_heading":
        target = None
        for p in original_paragraphs:
            sty = p.style.name if p.style else ""
            if sty.startswith(("Title", "Heading")) and (anchor or "") in p.text:
                target = p
                break
        if target is None:
            raise ValueError(f"no heading containing {anchor!r} found")
        target_el = target._element
        for el in reversed(new_elements):
            target_el.addnext(el)
    elif position == "after_paragraph":
        idx = _resolve_paragraph_index(original_paragraphs, anchor)
        anchor_para_el = original_paragraphs[idx]._element
        # Skip past any tables that immediately follow the anchor (caption
        # + its table form a semantic unit; "after the caption" usually
        # means "after the caption-and-table unit").
        target_el = _advance_past_following_tables(anchor_para_el)
        for el in reversed(new_elements):
            target_el.addnext(el)
    elif position == "before_paragraph":
        idx = _resolve_paragraph_index(original_paragraphs, anchor)
        target_el = original_paragraphs[idx]._element
        for el in new_elements:
            target_el.addprevious(el)
    elif position == "after_section":
        # Insert at the end of the section identified by ``anchor`` heading
        # — i.e. just before the next paragraph styled ``Title`` or
        # ``Heading N`` of same/higher level. Falls back to appending at
        # body end (before sectPr) when there's no follow-up heading.
        _, _, content_end = _resolve_section_bounds(original_paragraphs, anchor)
        if content_end < len(original_paragraphs):
            target_el = original_paragraphs[content_end]._element
            for el in new_elements:
                target_el.addprevious(el)
        else:
            section_pr = doc.element.body.find(qn("w:sectPr"))
            if section_pr is not None:
                for el in new_elements:
                    section_pr.addprevious(el)
            # else: new_elements were already appended at end → leave them.

    # Telemetry: count by block kind
    n_paragraphs, n_tables = _count_blocks(new_elements)
    return {
        "inserted_paragraphs": n_paragraphs,
        "inserted_tables": n_tables,
        "position": position,
        "anchor": anchor,
        "format": "markdown",
    }


def _resolve_paragraph_index(body: list[Any], anchor: Any) -> int:
    """Coerce an ``anchor`` (int, "3", or text substring) to a body index."""
    if isinstance(anchor, int):
        idx = anchor
    elif isinstance(anchor, str) and anchor.lstrip("-").isdigit():
        idx = int(anchor)
    elif isinstance(anchor, str):
        # Substring match on paragraph text (first hit).
        for i, p in enumerate(body):
            if anchor in p.text:
                return i
        raise ValueError(f"no paragraph containing {anchor!r} found in body")
    else:
        raise ValueError(f"anchor must be int or str, got {type(anchor).__name__}")
    if not (0 <= idx < len(body)):
        raise ValueError(
            f"paragraph index {idx} out of range (document has {len(body)} body paragraphs)"
        )
    return idx


def _heading_level_of_anchor(doc, anchor_text: str) -> Optional[int]:
    """Return the heading level (0=Title, 1-6=Heading 1..6) of the first body
    paragraph whose text contains ``anchor_text`` AND is styled as a heading.
    Returns None if no matching heading paragraph is found.
    """
    for p in _iter_body_paragraphs(doc):
        sty = p.style.name if p.style else ""
        if sty == "Title" and anchor_text in p.text:
            return 0
        if sty.startswith("Heading ") and anchor_text in p.text:
            try:
                return int(sty.split()[1])
            except (IndexError, ValueError):
                return None
    return None


def _advance_past_following_tables(target_el):
    """Given a ``<w:p>`` body element, walk its forward siblings and return
    the LAST element of any contiguous run of ``<w:tbl>`` elements directly
    following it. If no tables follow, return ``target_el`` unchanged.

    Why: the model's natural mental model is that a caption paragraph "owns"
    the table directly beneath it. So when the user writes
    ``after_paragraph + anchor=<caption text>``, they almost always mean
    "after the caption + its table", NOT "between the caption and its
    table". OOXML places them as separate body siblings, so without this
    helper a literal addnext would wedge new content into that gap, leaving
    the caption stranded above the new content and the table below.
    """
    from docx.oxml.ns import qn

    tbl_tag = qn("w:tbl")
    cur = target_el
    while True:
        nxt = cur.getnext()
        if nxt is None or nxt.tag != tbl_tag:
            break
        cur = nxt
    return cur


_BODY_HEADING_PREFIXES = ("Title", "Heading", "Subtitle", "TOC ")


def _resolve_target_paragraphs(
    doc,
    *,
    paragraph_index: Optional[int] = None,
    paragraph_indexes: Optional[list[int]] = None,
    anchor: Optional[str] = None,
    style_filter: Any = None,  # str | list[str] | None
) -> list[tuple[int, Any]]:
    """Resolve a multi-targeting selector into a list of ``(index, paragraph)``.

    Matches body paragraphs only (consistent with reader.get_outline).
    Exactly one selector should be provided.

    ``style_filter`` accepts:
        - a string: substring match against ``para.style.name`` (e.g.
          ``"Heading"`` matches H1-H6, ``"Heading 1"`` matches H1 only).
        - a list of strings: any-match — paragraph included if ANY of the
          substrings matches its style name. Use this to cover real-world
          docs that mix multiple body styles like ``["Normal", "Body Text",
          "FirstParagraph"]``.
        - the special string ``"!Heading"``: matches paragraphs whose style
          is NOT a heading / title / TOC entry (i.e. body content). Use this
          to apply uniform body formatting regardless of which body style
          the document was authored with.
    """
    body = _iter_body_paragraphs(doc)
    selectors_set = sum(
        1 for v in (paragraph_index, paragraph_indexes, anchor, style_filter)
        if v is not None
    )
    if selectors_set == 0:
        raise ValueError(
            "specify at least one of paragraph_index / paragraph_indexes / "
            "anchor / style_filter"
        )
    if selectors_set > 1:
        raise ValueError(
            "specify exactly one of paragraph_index / paragraph_indexes / "
            "anchor / style_filter"
        )

    if paragraph_index is not None:
        if not (0 <= paragraph_index < len(body)):
            raise ValueError(
                f"paragraph_index {paragraph_index} out of range "
                f"(document has {len(body)} body paragraphs)"
            )
        return [(paragraph_index, body[paragraph_index])]

    if paragraph_indexes is not None:
        result: list[tuple[int, Any]] = []
        for idx in paragraph_indexes:
            if not (0 <= idx < len(body)):
                raise ValueError(
                    f"paragraph_indexes contains {idx} out of range "
                    f"(document has {len(body)} body paragraphs)"
                )
            result.append((idx, body[idx]))
        return result

    if anchor is not None:
        return [(i, p) for i, p in enumerate(body) if anchor in p.text]

    # style_filter — string, list[str], or sentinel "!Heading"
    def _style_name(p) -> str:
        return p.style.name if p.style else ""

    if isinstance(style_filter, str):
        if style_filter == "!Heading":
            return [
                (i, p) for i, p in enumerate(body)
                if not _style_name(p).startswith(_BODY_HEADING_PREFIXES)
            ]
        return [(i, p) for i, p in enumerate(body) if style_filter in _style_name(p)]

    if isinstance(style_filter, (list, tuple)):
        filters = list(style_filter)
        return [
            (i, p) for i, p in enumerate(body)
            if any(f in _style_name(p) for f in filters)
        ]

    raise ValueError(
        f"style_filter must be str or list[str] (got {type(style_filter).__name__})"
    )


def _set_run_font_all_slots(run, font_name: str) -> None:
    """Set ascii + hAnsi + eastAsia font slots on a run.

    python-docx's ``run.font.name = X`` only writes the ascii slot, leaving
    eastAsia at its default. For CJK documents this means setting font_name
    has NO visible effect on Chinese characters — the latin slot changes
    but the CJK characters keep rendering in the eastAsia font (e.g. the
    document's default 方正仿宋简体). This helper writes all three slots so
    Chinese AND latin characters both pick up the new face.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    run.font.name = font_name  # keep python-docx accessor consistent


def _para_effective_font_size_pt(para) -> Optional[float]:
    """Best-effort read of a paragraph's effective font size in points.
    Looks at the paragraph's first run, then the paragraph's style chain,
    then the document's default. Returns None if nothing resolves.
    """
    for run in para.runs:
        sz = run.font.size
        if sz is not None:
            return sz.pt
    style = para.style
    while style is not None:
        sz = getattr(style.font, "size", None)
        if sz is not None:
            return sz.pt
        style = getattr(style, "base_style", None)
    return None


def _apply_format_to_paragraph(
    para,
    *,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    underline: Optional[bool] = None,
    font_size: Optional[int] = None,
    font_name: Optional[str] = None,
    color_hex: Optional[str] = None,
    line_spacing: Optional[float] = None,
    first_line_indent_chars: Optional[float] = None,
    first_line_indent_pt: Optional[float] = None,
    space_before_pt: Optional[float] = None,
    space_after_pt: Optional[float] = None,
) -> int:
    """Apply run-level formatting + paragraph-level layout to one paragraph.
    Returns the number of runs touched (paragraph-level fields don't count).

    ``line_spacing``: float multiplier (e.g. 1.5 for 1.5× spacing). Sets
    python-docx's ``paragraph_format.line_spacing``.

    ``first_line_indent_chars``: Chinese-style 段首空 N 格. Converted to
    points using the paragraph's effective font size (or ``font_size`` if
    that's part of this same call). Mutually exclusive with
    ``first_line_indent_pt`` — if both are passed, ``_pt`` wins.

    ``space_before_pt`` / ``space_after_pt``: paragraph spacing before/after,
    in points.
    """
    from docx.shared import Pt, RGBColor

    color_rgb = None
    if color_hex:
        clean = color_hex.lstrip("#")
        if len(clean) != 6:
            raise ValueError(f"color_hex must be 6 hex chars (got {color_hex!r})")
        try:
            color_rgb = RGBColor.from_string(clean)
        except Exception as e:
            raise ValueError(f"invalid color_hex {color_hex!r}: {e}") from e

    touched = 0
    for run in para.runs:
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic
        if underline is not None:
            run.underline = underline
        if font_size is not None:
            run.font.size = Pt(int(font_size))
        if font_name:
            _set_run_font_all_slots(run, font_name)
        if color_rgb is not None:
            run.font.color.rgb = color_rgb
        touched += 1

    # Paragraph-level layout.
    pf = para.paragraph_format
    if line_spacing is not None:
        pf.line_spacing = float(line_spacing)
    if first_line_indent_pt is not None:
        pf.first_line_indent = Pt(float(first_line_indent_pt))
    elif first_line_indent_chars is not None:
        # Convert N characters → points using this paragraph's effective font
        # size. CJK characters are full-width, so 1 char ≈ 1× font size in
        # points (the standard 公文 convention).
        size_pt = (
            float(font_size) if font_size is not None
            else _para_effective_font_size_pt(para) or 12.0
        )
        pf.first_line_indent = Pt(size_pt * float(first_line_indent_chars))
    if space_before_pt is not None:
        pf.space_before = Pt(float(space_before_pt))
    if space_after_pt is not None:
        pf.space_after = Pt(float(space_after_pt))

    return touched


_FORMAT_FIELDS = (
    "bold", "italic", "underline", "font_size", "font_name", "color_hex",
    "line_spacing", "first_line_indent_chars", "first_line_indent_pt",
    "space_before_pt", "space_after_pt",
)


def _op_format_text(
    doc,
    *,
    paragraph_index: Optional[int] = None,
    paragraph_indexes: Optional[list[int]] = None,
    anchor: Optional[str] = None,
    style_filter: Optional[str] = None,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    underline: Optional[bool] = None,
    font_size: Optional[int] = None,
    font_name: Optional[str] = None,
    color_hex: Optional[str] = None,
    line_spacing: Optional[float] = None,
    first_line_indent_chars: Optional[float] = None,
    first_line_indent_pt: Optional[float] = None,
    space_before_pt: Optional[float] = None,
    space_after_pt: Optional[float] = None,
) -> dict[str, Any]:
    fields = {
        "bold": bold, "italic": italic, "underline": underline,
        "font_size": font_size, "font_name": font_name, "color_hex": color_hex,
        "line_spacing": line_spacing,
        "first_line_indent_chars": first_line_indent_chars,
        "first_line_indent_pt": first_line_indent_pt,
        "space_before_pt": space_before_pt, "space_after_pt": space_after_pt,
    }
    if all(v is None for v in fields.values()):
        raise ValueError(
            f"at least one formatting field must be set: {sorted(fields)}"
        )

    targets = _resolve_target_paragraphs(
        doc,
        paragraph_index=paragraph_index,
        paragraph_indexes=paragraph_indexes,
        anchor=anchor,
        style_filter=style_filter,
    )
    if not targets:
        return {
            "paragraphs_touched": 0,
            "runs_touched": 0,
            "warning": "no paragraphs matched the selector",
        }

    runs_total = 0
    for _, para in targets:
        runs_total += _apply_format_to_paragraph(para, **fields)

    return {
        "paragraphs_touched": len(targets),
        "runs_touched": runs_total,
        "indexes_touched": [i for i, _ in targets],
        "applied": {k: v for k, v in fields.items() if v is not None},
    }


def set_cell_multiline(cell, value: str, font: Optional[str] = None) -> None:
    """Replace a cell's content with one paragraph per ``\\n``-delimited line.

    Critical: ``cell.text = "...\\n..."`` puts a literal LF inside the cell's
    single ``<w:t>`` element. Word may render it as a tofu glyph; LibreOffice
    headless → PDF conversion typically renders it as 乱码 (boxes / wrong chars).
    The fix is multi-paragraph: each ``\\n``-split line becomes its own
    ``<w:p>`` inside the cell, preserving font/style.

    Empty input is permitted — produces a single empty paragraph (cells must
    have ≥ 1 paragraph per OOXML).
    """
    from .styles import apply_cjk_font_to_para

    lines = value.split("\n") if value else [""]

    # Drop all existing paragraphs except the first (cells require ≥ 1).
    paras = list(cell.paragraphs)
    for extra in paras[1:]:
        extra._element.getparent().remove(extra._element)

    if not paras:
        first = cell.add_paragraph(lines[0])
    else:
        first = paras[0]
        # Reset first paragraph: keep one run, set its text, drop the rest.
        for r in list(first.runs)[1:]:
            r._element.getparent().remove(r._element)
        if first.runs:
            first.runs[0].text = lines[0]
        else:
            first.add_run(lines[0])
    if font:
        apply_cjk_font_to_para(first, font)

    for line in lines[1:]:
        new_p = cell.add_paragraph(line)
        if font:
            apply_cjk_font_to_para(new_p, font)


def _op_replace_paragraph(
    doc,
    *,
    anchor: Any,
    new_text: str,
    style: Optional[str] = None,
    format: str = "auto",
) -> dict[str, Any]:
    """Replace ONE body paragraph (resolved by anchor/index) with new content.

    ``format`` (``"auto"`` default / ``"markdown"`` / ``"text"``) decides how
    ``new_text`` is rendered — see ``_resolve_content_format``. In ``markdown``
    mode the new_text is parsed into proper Word blocks (headings / lists /
    bold) and ``style`` is ignored (block styles come from the markdown
    structure); ``auto`` does this automatically when ``new_text`` looks like
    markdown. ``auto`` on plain prose behaves exactly like the legacy path.

    Plain-text path: multi-paragraph ``new_text`` (containing ``\\n``) is
    auto-split: the first line goes into the original paragraph (preserving its
    position + style), subsequent lines become new paragraphs inserted right
    after. This avoids the silent failure mode where ``\\n`` ends up as a
    literal LF in a single ``<w:t>``. If ``style`` is provided, ALL produced
    paragraphs adopt it; otherwise the original paragraph's style is kept.
    """
    from .styles import apply_cjk_font_to_para, font_for_style

    body = _iter_body_paragraphs(doc)
    idx = _resolve_paragraph_index(body, anchor)
    target = body[idx]
    original_style = target.style.name if target.style else None
    final_style = style if style is not None else original_style

    # Snapshot the target's body formatting (indent / spacing / justify) so the
    # replacement keeps looking like the surrounding text instead of reverting
    # to bare Normal-style paragraphs.
    ref_fmt = _capture_body_format(target)

    # ── Markdown path — render blocks where the target paragraph sat ──
    # Emit BEFORE removing the target so a parse error (e.g. table syntax)
    # leaves the document untouched.
    if _resolve_content_format(new_text, format) == "markdown":
        new_elements = _emit_markdown_blocks_safe(doc, new_text)
        _relocate_after(target._element, new_elements)
        parent = target._element.getparent()
        if parent is not None:
            parent.remove(target._element)
        _apply_body_format_to_blocks(new_elements, ref_fmt)
        n_p, n_tbl = _count_blocks(new_elements)
        result: dict[str, Any] = {
            "replaced_index": idx,
            "format": "markdown",
            "new_paragraph_count": n_p,
            "new_table_count": n_tbl,
            "new_chars": len(new_text or ""),
        }
        if style is not None:
            result["style_ignored"] = style  # markdown blocks carry own styles
        return result

    lines = _split_body_lines(new_text)
    para_font = font_for_style(final_style)

    # First line into target paragraph (keeps position + its own indent/spacing).
    for run in list(target.runs)[1:]:
        run._element.getparent().remove(run._element)
    if target.runs:
        target.runs[0].text = lines[0]
    else:
        target.add_run(lines[0])
    if final_style:
        try:
            target.style = doc.styles[final_style]
        except KeyError:
            pass  # unknown style → leave as-is
    apply_cjk_font_to_para(target, para_font)

    # Remaining lines → new paragraphs inserted in document order right after
    # target, each inheriting the target's body format so they don't lose the
    # first-line indent / line spacing the original paragraph had.
    cur_el = target._element
    for line in lines[1:]:
        new_p = doc.add_paragraph(line)
        if final_style:
            try:
                new_p.style = doc.styles[final_style]
            except KeyError:
                pass
        _apply_body_format(new_p, ref_fmt)
        apply_cjk_font_to_para(new_p, para_font)
        # Move from end-of-document to right after cur_el.
        cur_el.addnext(new_p._element)
        cur_el = new_p._element

    result = {
        "replaced_index": idx,
        "format": "text",
        "style": final_style,
        "new_paragraph_count": len(lines),
        "new_chars": len(new_text),
    }
    if _looks_like_markdown(new_text or ""):
        result["warning"] = _MARKDOWN_LITERAL_WARNING
    return result


def _op_delete_paragraph(
    doc,
    *,
    anchor: Any,
) -> dict[str, Any]:
    """Remove ONE paragraph from the body. Resolves anchor like insert_text."""
    body = _iter_body_paragraphs(doc)
    idx = _resolve_paragraph_index(body, anchor)
    target = body[idx]
    parent = target._element.getparent()
    if parent is None:
        raise ValueError("paragraph has no parent — already detached?")
    parent.remove(target._element)
    return {"deleted_index": idx}


def _op_set_cell_text(
    doc,
    *,
    table_index: int,
    row: int,
    col: int,
    text: str,
    preserve_format: bool = True,
) -> dict[str, Any]:
    """Set the text of a single table cell. Multi-line text (``\\n``) becomes
    multiple paragraphs inside the cell so the line breaks survive PDF export.
    """
    from .styles import BODY_FONT

    if not (0 <= table_index < len(doc.tables)):
        raise ValueError(
            f"table_index {table_index} out of range (document has {len(doc.tables)} tables)"
        )
    table = doc.tables[table_index]
    if not (0 <= row < len(table.rows)):
        raise ValueError(f"row {row} out of range (table has {len(table.rows)} rows)")
    if not (0 <= col < len(table.rows[row].cells)):
        raise ValueError(
            f"col {col} out of range (row has {len(table.rows[row].cells)} cells)"
        )

    cell = table.rows[row].cells[col]
    # ``preserve_format`` controls whether a CJK font is re-applied. Either
    # way we use the multi-paragraph pathway so ``\\n`` doesn't bleed into
    # ``<w:t>`` as a literal LF (renders as 乱码 in LibreOffice → PDF).
    set_cell_multiline(cell, text, font=BODY_FONT if preserve_format else None)

    line_count = len(text.split("\n")) if text else 1
    return {
        "table_index": table_index, "row": row, "col": col,
        "preserve_format": preserve_format,
        "line_count": line_count,
    }


def _op_fill_table(
    doc,
    *,
    table_index: int,
    rows: list[list[str]],
    mode: str = "append",
    has_header: bool = True,
) -> dict[str, Any]:
    """Append or overwrite data rows in an existing table.

    Modes:
        - "append": add rows after existing rows.
        - "overwrite": keep header row (if has_header), drop all data rows,
          then add the new rows.

    Cell formatting is set via ``cell.text = value``, which preserves the
    table style but replaces cell content. Maintains a cell count compatible
    with the table's existing column count; extra values are ignored, missing
    values are left blank.
    """
    if mode not in ("append", "overwrite"):
        raise ValueError(f"mode must be 'append' or 'overwrite', got {mode!r}")
    if not (0 <= table_index < len(doc.tables)):
        raise ValueError(
            f"table_index {table_index} out of range (document has {len(doc.tables)} tables)"
        )
    if not isinstance(rows, list) or not rows:
        raise ValueError("'rows' must be a non-empty list of cell lists")

    table = doc.tables[table_index]
    existing = list(table.rows)
    if not existing:
        raise ValueError("target table has no rows; cannot infer column layout")
    col_count = len(existing[0].cells)

    if mode == "overwrite":
        keep = 1 if has_header else 0
        for r in existing[keep:]:
            r._element.getparent().remove(r._element)

    from .styles import BODY_FONT

    added = 0
    for row_vals in rows:
        new_row = table.add_row()
        for ci, cell in enumerate(new_row.cells[:col_count]):
            raw = row_vals[ci] if ci < len(row_vals) else ""
            value = "" if raw is None else str(raw)
            # Multi-paragraph cell text so ``\\n`` survives PDF conversion.
            set_cell_multiline(cell, value, font=BODY_FONT)
        added += 1

    return {
        "table_index": table_index,
        "rows_added": added,
        "mode": mode,
        "column_count": col_count,
    }


# ── public single-shot API (open → op → save → return) ──────────────────────


def _open(input_filename: str):
    from docx import Document
    return Document(str(input_path(input_filename)))


def _save(doc, output_filename: str) -> None:
    out = output_path(output_filename)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))


def search_and_replace(
    *,
    input_filename: str,
    output_filename: str,
    find: str,
    replace: str,
    scope: str = "all",
    regex: bool = False,
    lenient: bool = False,
) -> dict[str, Any]:
    doc = _open(input_filename)
    out = _op_search_and_replace(
        doc, find=find, replace=replace, scope=scope, regex=regex, lenient=lenient,
    )
    _save(doc, output_filename)
    return {"output_filename": output_filename, **out}


def replace_many(
    *,
    input_filename: str,
    output_filename: str,
    replacements: list[dict[str, Any]],
) -> dict[str, Any]:
    doc = _open(input_filename)
    out = _op_replace_many(doc, replacements=replacements)
    _save(doc, output_filename)
    return {"output_filename": output_filename, **out}


def fill_placeholders(
    *,
    input_filename: str,
    output_filename: str,
    mapping: dict[str, str],
    pattern: str = r"\{\{(\w+)\}\}",
) -> dict[str, Any]:
    doc = _open(input_filename)
    out = _op_fill_placeholders(doc, mapping=mapping, pattern=pattern)
    _save(doc, output_filename)
    return {"output_filename": output_filename, **out}


def list_placeholders(
    *,
    input_filename: str,
    pattern: str = r"\{\{(\w+)\}\}",
) -> dict[str, Any]:
    # Read-only — no save.
    doc = _open(input_filename)
    return _op_list_placeholders(doc, pattern=pattern)


def insert_text(
    *,
    input_filename: str,
    output_filename: str,
    text: str,
    position: str = "end",
    anchor: Optional[str] = None,
    style: Optional[str] = None,
    style_for_all: bool = False,
    format: str = "auto",
) -> dict[str, Any]:
    doc = _open(input_filename)
    out = _op_insert_text(
        doc, text=text, position=position, anchor=anchor,
        style=style, style_for_all=style_for_all, format=format,
    )
    _save(doc, output_filename)
    return {"output_filename": output_filename, **out}


def format_text(
    *,
    input_filename: str,
    output_filename: str,
    paragraph_index: Optional[int] = None,
    paragraph_indexes: Optional[list[int]] = None,
    anchor: Optional[str] = None,
    style_filter: Optional[str] = None,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    underline: Optional[bool] = None,
    font_size: Optional[int] = None,
    font_name: Optional[str] = None,
    color_hex: Optional[str] = None,
    line_spacing: Optional[float] = None,
    first_line_indent_chars: Optional[float] = None,
    first_line_indent_pt: Optional[float] = None,
    space_before_pt: Optional[float] = None,
    space_after_pt: Optional[float] = None,
) -> dict[str, Any]:
    doc = _open(input_filename)
    out = _op_format_text(
        doc,
        paragraph_index=paragraph_index,
        paragraph_indexes=paragraph_indexes,
        anchor=anchor,
        style_filter=style_filter,
        bold=bold, italic=italic, underline=underline,
        font_size=font_size, font_name=font_name, color_hex=color_hex,
        line_spacing=line_spacing,
        first_line_indent_chars=first_line_indent_chars,
        first_line_indent_pt=first_line_indent_pt,
        space_before_pt=space_before_pt, space_after_pt=space_after_pt,
    )
    _save(doc, output_filename)
    return {"output_filename": output_filename, **out}



# ── batch / atomic editing ──────────────────────────────────────────────────


def _resolve_section_bounds(body: list[Any], heading_anchor: str) -> tuple[int, int, int]:
    """Return ``(heading_idx, content_start, content_end_exclusive)`` for the
    section identified by ``heading_anchor``.

    ``heading_idx`` is the index of the matching heading paragraph; the section's
    body content lives at ``[content_start, content_end_exclusive)`` and ends
    just before the next heading whose level is ≤ the matched one (or at end of
    document if there is none).
    """
    heading_idx = -1
    heading_level = -1
    for i, p in enumerate(body):
        sty = p.style.name if p.style else ""
        if not (sty == "Title" or sty.startswith("Heading ")):
            continue
        if heading_anchor in p.text:
            heading_idx = i
            if sty == "Title":
                heading_level = 0
            else:
                try:
                    heading_level = int(sty.split()[1])
                except (IndexError, ValueError):
                    heading_level = 99  # treat as deepest
            break
    if heading_idx < 0:
        raise ValueError(f"no heading containing {heading_anchor!r} found")

    content_end = len(body)
    for j in range(heading_idx + 1, len(body)):
        sty = body[j].style.name if body[j].style else ""
        if sty == "Title":
            level = 0
        elif sty.startswith("Heading "):
            try:
                level = int(sty.split()[1])
            except (IndexError, ValueError):
                continue
        else:
            continue
        if level <= heading_level:
            content_end = j
            break

    return heading_idx, heading_idx + 1, content_end


def _op_replace_section(
    doc,
    *,
    heading_anchor: str,
    new_content: str,
    preserve_heading: bool = True,
    style: Optional[str] = None,
    format: str = "auto",
) -> dict[str, Any]:
    """Replace an entire heading section's content with ``new_content``.

    A "section" is the heading paragraph + every following paragraph until the
    next paragraph styled ``Title`` or ``Heading N`` of the SAME or HIGHER level
    (smaller N means higher level). With ``preserve_heading=True`` (default),
    the heading itself is kept and only the body content is replaced; with
    ``False``, the heading is also replaced (using its first ``\\n``-split line
    as the new heading text).

    ``format`` (``"auto"`` default / ``"markdown"`` / ``"text"``) decides how
    ``new_content`` is rendered — see ``_resolve_content_format``. Whole-section
    rewrites are exactly where the LLM dumps a Markdown draft, so ``auto``
    renders ``###`` headings / ``- `` lists / ``**bold**`` as real Word
    formatting when the content looks like markdown, and ``style`` is ignored in
    that case (block styles come from the markdown). With ``preserve_heading=
    True`` the markdown ``new_content`` is the section BODY only — don't repeat
    the section heading or it renders twice.

    Plain-text path: ``new_content`` is split on ``\\n`` into paragraphs;
    BLANK lines act as separators only (no empty paragraphs left behind — a
    markdown-style ``\\n\\n`` between paragraphs is the norm). New paragraphs
    inherit the section body's first-line indent / line spacing so the rewrite
    matches the surrounding document instead of reverting to bare Normal.

    USE WHEN:
        Rewriting a whole section ("把 6.2 整段换掉")。比串联 N 次
        replace_paragraph / delete_paragraph 安全得多——单次原子操作，无索引漂移。
    """
    from .styles import apply_cjk_font_to_para, font_for_style

    body = _iter_body_paragraphs(doc)
    h_idx, content_start, content_end = _resolve_section_bounds(body, heading_anchor)
    heading_para = body[h_idx]
    para_font = font_for_style(style) if style else None

    # Snapshot the section body's formatting (first removed body paragraph, else
    # the document's prevailing indented body paragraph) so the rewrite keeps the
    # same first-line indent / line spacing instead of reverting to bare Normal.
    ref_fmt: list[Any] = []
    for p in body[content_start:content_end]:
        if not _is_body_block(p._element):
            continue
        ref_fmt = _capture_body_format(p)
        if ref_fmt:
            break
    if not ref_fmt:
        _skip = {id(p._element) for p in body[content_start:content_end]}
        _skip.add(id(heading_para._element))
        ref_fmt = _scan_body_format(doc, skip=_skip)

    # ── Markdown path — render blocks in place of the section's content ──
    # Emit BEFORE deleting the old content so a parse error leaves the section
    # intact (no data loss).
    if _resolve_content_format(new_content, format) == "markdown":
        new_elements = _emit_markdown_blocks_safe(doc, new_content)
        if preserve_heading:
            for p in body[content_start:content_end]:
                parent = p._element.getparent()
                if parent is not None:
                    parent.remove(p._element)
            _relocate_after(heading_para._element, new_elements)
        else:
            # Drop heading + body; the markdown content (typically led by its
            # own heading) stands in for the whole section.
            _relocate_before(heading_para._element, new_elements)
            for p in body[content_start:content_end]:
                parent = p._element.getparent()
                if parent is not None:
                    parent.remove(p._element)
            h_parent = heading_para._element.getparent()
            if h_parent is not None:
                h_parent.remove(heading_para._element)
        _apply_body_format_to_blocks(new_elements, ref_fmt)
        n_p, n_tbl = _count_blocks(new_elements)
        result: dict[str, Any] = {
            "heading_index": h_idx,
            "format": "markdown",
            "removed_paragraphs": content_end - content_start,
            "new_paragraph_count": n_p,
            "new_table_count": n_tbl,
            "preserve_heading": preserve_heading,
        }
        if style is not None:
            result["style_ignored"] = style
        return result

    lines = _split_body_lines(new_content)

    # Resolve insertion point (the heading's element if preserving, else its
    # previous sibling — but with preserve_heading=False we re-use the heading
    # paragraph for the first new line, so insertion still anchors there).
    if preserve_heading:
        # Step 1: drop existing body content (start..end)
        for p in body[content_start:content_end]:
            parent = p._element.getparent()
            if parent is not None:
                parent.remove(p._element)
        # Step 2: insert new paragraphs right after the heading, each inheriting
        # the section body's original indent/spacing.
        cur_el = heading_para._element
        for line in lines:
            new_p = doc.add_paragraph(line)
            if style:
                try:
                    new_p.style = doc.styles[style]
                except KeyError:
                    pass
            _apply_body_format(new_p, ref_fmt)
            apply_cjk_font_to_para(new_p, para_font or font_for_style(None))
            cur_el.addnext(new_p._element)
            cur_el = new_p._element
        new_paragraph_count = len(lines)
    else:
        # preserve_heading=False — replace heading + body in one go. Reuse the
        # heading paragraph for the first new line (keeps position), drop the
        # rest of the section, then add subsequent lines after.
        for run in list(heading_para.runs)[1:]:
            run._element.getparent().remove(run._element)
        if heading_para.runs:
            heading_para.runs[0].text = lines[0]
        else:
            heading_para.add_run(lines[0])
        if style:
            try:
                heading_para.style = doc.styles[style]
            except KeyError:
                pass
        apply_cjk_font_to_para(heading_para, para_font or font_for_style(None))

        # Drop original body content under the (now-replaced) heading.
        for p in body[content_start:content_end]:
            parent = p._element.getparent()
            if parent is not None:
                parent.remove(p._element)

        cur_el = heading_para._element
        for line in lines[1:]:
            new_p = doc.add_paragraph(line)
            if style:
                try:
                    new_p.style = doc.styles[style]
                except KeyError:
                    pass
            _apply_body_format(new_p, ref_fmt)
            apply_cjk_font_to_para(new_p, para_font or font_for_style(None))
            cur_el.addnext(new_p._element)
            cur_el = new_p._element
        new_paragraph_count = len(lines)

    result = {
        "heading_index": h_idx,
        "format": "text",
        "removed_paragraphs": content_end - content_start,
        "new_paragraph_count": new_paragraph_count,
        "preserve_heading": preserve_heading,
    }
    if _looks_like_markdown(new_content or ""):
        result["warning"] = _MARKDOWN_LITERAL_WARNING
    return result


def _op_delete_range(
    doc,
    *,
    start_anchor: Any,
    end_anchor: Any,
    include_end: bool = False,
) -> dict[str, Any]:
    """Delete all body paragraphs from ``start_anchor`` (inclusive) to
    ``end_anchor`` (exclusive by default; inclusive when ``include_end=True``).

    Each anchor is resolved like ``delete_paragraph`` — int body index OR
    paragraph-text substring (first match). ``end_anchor`` must resolve to an
    index ≥ ``start_anchor``'s.

    USE WHEN:
        Removing a whole block atomically — much safer than a sequence of
        ``delete_paragraph`` ops with integer anchors (which suffer from
        index drift after each delete).
    """
    body = _iter_body_paragraphs(doc)
    start_idx = _resolve_paragraph_index(body, start_anchor)
    end_idx = _resolve_paragraph_index(body, end_anchor)
    if end_idx < start_idx:
        raise ValueError(
            f"end_anchor resolved to index {end_idx} which is before start "
            f"index {start_idx}"
        )
    upper = end_idx + 1 if include_end else end_idx
    removed = 0
    for p in body[start_idx:upper]:
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)
            removed += 1
    return {
        "start_index": start_idx,
        "end_index": end_idx,
        "include_end": include_end,
        "removed_paragraphs": removed,
    }


def _op_move_table(
    doc,
    *,
    table_index: int,
    position: str,
    anchor: Any = None,
) -> dict[str, Any]:
    """Move an EXISTING table to a new position. Atomic single op.

    Use this when the user says "move the table at the end up to section X" or
    "put this table after the conclusion heading". Far cheaper than the
    "delete table content + add_table at new location" dance — and unlike that
    dance it preserves all the table's original styling, merged cells, etc.

    Args:
        table_index: 0-based index of the table to move (as listed in
            ``get_outline.tables``).
        position: ``"end"`` / ``"start"`` / ``"after_heading"`` /
            ``"after_paragraph"`` / ``"before_paragraph"``.
        anchor: required for non-``end``/``start`` positions; heading text
            substring or paragraph index/substring.

    Returns:
        ``{"table_index", "moved_from_preceding_index", "position", "anchor"}``.
    """
    from .builder import (
        _VALID_TABLE_POSITIONS,
        _resolve_anchor_paragraph,
    )

    if position not in _VALID_TABLE_POSITIONS:
        raise ValueError(
            f"position must be one of {_VALID_TABLE_POSITIONS}, got {position!r}"
        )
    if position in ("after_heading", "after_paragraph", "before_paragraph") and anchor is None:
        raise ValueError(f"position={position!r} requires 'anchor'")
    if not (0 <= table_index < len(doc.tables)):
        raise ValueError(
            f"table_index {table_index} out of range (document has {len(doc.tables)} tables)"
        )

    # Capture original spatial position for telemetry.
    table = doc.tables[table_index]
    table_el = table._element
    body = doc.element.body
    from docx.oxml.ns import qn
    p_tag = qn("w:p")
    body_p_count = 0
    moved_from_preceding = None
    for child in body:
        if child is table_el:
            moved_from_preceding = body_p_count - 1 if body_p_count > 0 else None
            break
        if child.tag == p_tag:
            body_p_count += 1

    # Resolve target paragraph BEFORE detaching the table (otherwise indexes
    # could shift if the table happens to sit between target candidates).
    target_el = None
    if position not in ("end", "start"):
        target_el = _resolve_anchor_paragraph(doc, position, anchor)

    # Detach + re-insert. lxml's addnext / addprevious handle re-parenting
    # correctly even when the element is currently elsewhere in the same tree.
    if position == "end":
        # Move to end of body, before sectPr if present.
        from docx.oxml.ns import qn as _qn
        sect = body.find(_qn("w:sectPr"))
        if sect is not None:
            sect.addprevious(table_el)
        else:
            body.append(table_el)
    elif position == "start":
        if len(body) > 0:
            body[0].addprevious(table_el)
        else:
            body.append(table_el)
    elif position in ("after_heading", "after_paragraph"):
        target_el.addnext(table_el)
    elif position == "before_paragraph":
        target_el.addprevious(table_el)

    return {
        "table_index": table_index,
        "moved_from_preceding_index": moved_from_preceding,
        "position": position,
        "anchor": anchor,
    }


def _op_add_table(
    doc,
    *,
    rows: list[list[str]] | None = None,
    markdown: str | None = None,
    has_header: bool = True,
    caption: Optional[str] = None,
    position: str,
    anchor: Any = None,
    auto_merge_empty: bool = True,
) -> dict[str, Any]:
    """Apply-edits op: insert a styled table at a specific position in the document.

    Thin shim over ``builder._add_table_to_doc`` so ``apply_edits`` can drop
    a table mid-document as part of a larger atomic edit batch.

    ``position`` is REQUIRED — no default. See ``builder._add_table_to_doc``
    for accepted values.
    """
    from .builder import _add_table_to_doc
    return _add_table_to_doc(
        doc,
        rows=rows, markdown=markdown,
        has_header=has_header, caption=caption,
        position=position, anchor=anchor,
        auto_merge_empty=auto_merge_empty,
    )


# python-docx attribute → OOXML core property → user-facing CLI name.
# Mirrors the .NET CLI's `edit update-field` which writes the same
# ``dc:title`` / ``dc:creator`` / etc. core properties.
_DOCPROP_FIELD_MAP = {
    "TITLE":       "title",
    "AUTHOR":      "author",
    "SUBJECT":     "subject",
    "KEYWORDS":    "keywords",
    "DESCRIPTION": "comments",   # OOXML dc:description ↔ python-docx .comments
    "CATEGORY":    "category",
}


def _op_update_field(doc, *, field: str, value: str) -> dict[str, Any]:
    """Apply-edits op: update a document core property (TITLE/AUTHOR/etc).

    Updates the same OOXML metadata fields the .NET CLI's ``edit update-field``
    targets — ``docProps/core.xml`` (``dc:title``, ``dc:creator``,
    ``dc:subject``, ``cp:keywords``, ``dc:description``, ``cp:category``).

    Args:
        field: case-insensitive field name. One of ``"TITLE"``, ``"AUTHOR"``,
               ``"SUBJECT"``, ``"KEYWORDS"``, ``"DESCRIPTION"``, ``"CATEGORY"``.
        value: new value for the field. Strings only.

    Raises:
        ValueError: unknown field, or value is not a string.
    """
    if not isinstance(field, str) or not field:
        raise ValueError("'field' must be a non-empty string")
    if not isinstance(value, str):
        raise ValueError(
            f"'value' must be a string, got {type(value).__name__}"
        )
    attr = _DOCPROP_FIELD_MAP.get(field.upper())
    if attr is None:
        raise ValueError(
            f"unknown field {field!r}; "
            f"supported: {sorted(_DOCPROP_FIELD_MAP)}"
        )
    setattr(doc.core_properties, attr, value)
    return {"field": field.upper(), "value": value, "attr": attr}


_ALIGNMENT_MAP = {
    "left": "LEFT",
    "center": "CENTER",
    "centre": "CENTER",
    "right": "RIGHT",
}


def _op_insert_image(
    doc,
    *,
    image_path: Optional[str] = None,
    image_filename: Optional[str] = None,
    position: str = "end",
    anchor: Optional[str] = None,
    width_inches: Optional[float] = None,
    width_cm: Optional[float] = None,
    alignment: Optional[str] = None,
) -> dict[str, Any]:
    """Apply-edits op: insert an image into the document at a positioned anchor.

    Image source — pass ONE of:
      * ``image_path`` (**preferred**): an absolute sandbox path to the image,
        e.g. ``/workspace/chart1.png``. This is the simple path: first stage the
        image into the sandbox (``sandbox_put_artifact(artifact_id=<chart
        file_id>, dest_path="/workspace/chart1.png")``), then reference that path
        here directly — no ``--image`` flag, no alias.
      * ``image_filename`` (legacy): a workdir-relative name. The skill's
        ``apply_edits.py`` resolves an ``image_file_id`` → ``--image`` mapping →
        workdir file → this name before dispatch.

    Position semantics mirror ``_op_insert_text`` — supports ``end`` / ``start``
    / ``after_heading`` / ``after_section`` / ``after_paragraph`` /
    ``before_paragraph``. Inserting into table cells is **not** supported by
    this op (use a separate workflow for that).

    Width: pass ``width_inches`` OR ``width_cm`` to constrain the image width
    (height auto-scales). Both omitted → image is inserted at native size,
    which often overflows page margins; recommend the LLM pass an explicit
    width for any non-trivial image.

    Alignment: ``"left" | "center" | "right"`` — applied to the new paragraph.
    Default leaves the paragraph default (typically left).
    """
    from pathlib import Path

    from docx.shared import Cm, Inches

    from ._handle import input_path
    from .styles import (
        ANCHOR_REQUIRED_POSITIONS as _ANCHOR_REQUIRED,
        INSERT_POSITIONS as _VALID_POSITIONS,
    )

    if position not in _VALID_POSITIONS:
        raise ValueError(
            f"position must be one of {_VALID_POSITIONS}, got {position!r}"
        )
    if position in _ANCHOR_REQUIRED and anchor is None:
        raise ValueError(f"position={position!r} requires 'anchor'")
    if width_inches is not None and width_cm is not None:
        raise ValueError("pass only one of 'width_inches' or 'width_cm'")

    # Resolve the image source. ``image_path`` (a sandbox path) wins; an
    # absolute path is read directly off the real FS (the engine workdir is a
    # thread-local staging temp, so input_path() would NOT find /workspace/...).
    if image_path is not None:
        if not isinstance(image_path, str) or not image_path:
            raise ValueError("'image_path' must be a non-empty string")
        _p = Path(image_path)
        img_path = _p if _p.is_absolute() else input_path(image_path)
        if not img_path.is_file():
            raise FileNotFoundError(
                f"image not found at {img_path}; did you sandbox_put_artifact it "
                f"into the sandbox first?"
            )
    elif image_filename:
        img_path = input_path(image_filename)
    else:
        raise ValueError(
            "provide 'image_path' (sandbox path, e.g. /workspace/chart.png) "
            "or 'image_filename'"
        )

    width = None
    if width_inches is not None:
        if width_inches <= 0:
            raise ValueError("'width_inches' must be > 0")
        width = Inches(width_inches)
    elif width_cm is not None:
        if width_cm <= 0:
            raise ValueError("'width_cm' must be > 0")
        width = Cm(width_cm)

    align_enum = None
    if alignment is not None:
        key = alignment.strip().lower()
        if key not in _ALIGNMENT_MAP:
            raise ValueError(
                f"alignment must be one of {sorted(set(_ALIGNMENT_MAP))}, "
                f"got {alignment!r}"
            )
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        align_enum = getattr(WD_ALIGN_PARAGRAPH, _ALIGNMENT_MAP[key])

    body = _iter_body_paragraphs(doc)

    def _attach_picture(para) -> None:
        run = para.add_run()
        run.add_picture(str(img_path), width=width)
        if align_enum is not None:
            para.alignment = align_enum

    if position == "end":
        new_para = doc.add_paragraph()
        _attach_picture(new_para)
    elif position == "start":
        new_para = doc.add_paragraph()
        _attach_picture(new_para)
        if body:
            body[0]._element.addprevious(new_para._element)
    elif position == "after_heading":
        target = None
        for p in body:
            sty = p.style.name if p.style else ""
            if sty.startswith(("Title", "Heading")) and (anchor or "") in p.text:
                target = p
                break
        if target is None:
            raise ValueError(f"no heading containing {anchor!r} found")
        new_para = doc.add_paragraph()
        _attach_picture(new_para)
        target._element.addnext(new_para._element)
    elif position == "after_paragraph":
        idx = _resolve_paragraph_index(body, anchor)
        target = body[idx]
        anchor_el = _advance_past_following_tables(target._element)
        new_para = doc.add_paragraph()
        _attach_picture(new_para)
        anchor_el.addnext(new_para._element)
    elif position == "before_paragraph":
        idx = _resolve_paragraph_index(body, anchor)
        target = body[idx]
        new_para = doc.add_paragraph()
        _attach_picture(new_para)
        target._element.addprevious(new_para._element)
    elif position == "after_section":
        _, _, content_end = _resolve_section_bounds(body, anchor)
        new_para = doc.add_paragraph()
        _attach_picture(new_para)
        if content_end < len(body):
            body[content_end]._element.addprevious(new_para._element)
        else:
            from docx.oxml.ns import qn as _qn
            section_pr = doc.element.body.find(_qn("w:sectPr"))
            if section_pr is not None:
                section_pr.addprevious(new_para._element)

    width_repr: Optional[str] = None
    if width_inches is not None:
        width_repr = f"{width_inches}in"
    elif width_cm is not None:
        width_repr = f"{width_cm}cm"

    return {
        "image_source": image_path or image_filename,
        "position": position,
        "anchor": anchor,
        "width": width_repr,
        "alignment": alignment,
    }


# Map op name → primitive function. Each primitive takes (doc, **kwargs).
_OP_REGISTRY: dict[str, Callable[..., dict[str, Any]]] = {
    "replace": _op_search_and_replace,
    "replace_many": _op_replace_many,
    "fill_placeholders": _op_fill_placeholders,
    "insert": _op_insert_text,
    "insert_image": _op_insert_image,
    "format": _op_format_text,
    "replace_paragraph": _op_replace_paragraph,
    "delete_paragraph": _op_delete_paragraph,
    "delete_range": _op_delete_range,
    "set_cell_text": _op_set_cell_text,
    "fill_table": _op_fill_table,
    "add_table": _op_add_table,
    "move_table": _op_move_table,
    "replace_section": _op_replace_section,
    "update_field": _op_update_field,
}


def _normalize_int_anchor_deletes(ops: list[dict[str, Any]]) -> tuple[list[dict[str, Any]], int]:
    """Defend against index drift in chained ``delete_paragraph`` ops.

    When the LLM emits a run of consecutive ``delete_paragraph`` ops with
    integer anchors, executing them in input order is wrong: deleting paragraph
    85 shifts the original 86 down to position 85, so the next op's
    ``anchor=86`` actually deletes the original 87. Result: every other
    paragraph survives.

    The fix is to execute each contiguous run of int-anchored delete_paragraph
    ops in **descending** anchor order — deleting from the back keeps earlier
    indexes stable. Non-int anchors and other op types are left untouched
    (their semantics don't depend on neighboring deletions).

    Returns the rewritten op list AND the count of ops that were re-ordered
    (for telemetry).
    """
    out: list[dict[str, Any]] = []
    reordered = 0
    i = 0
    while i < len(ops):
        op = ops[i]
        if (
            isinstance(op, dict)
            and op.get("op") == "delete_paragraph"
            and isinstance(op.get("anchor"), int)
        ):
            # Collect the contiguous run
            run_start = i
            while (
                i < len(ops)
                and isinstance(ops[i], dict)
                and ops[i].get("op") == "delete_paragraph"
                and isinstance(ops[i].get("anchor"), int)
            ):
                i += 1
            run = ops[run_start:i]
            if len(run) > 1:
                run_sorted = sorted(run, key=lambda o: -o["anchor"])
                if run_sorted != run:
                    reordered += len(run)
                out.extend(run_sorted)
            else:
                out.extend(run)
        else:
            out.append(op)
            i += 1
    return out, reordered


def apply_edits(
    *,
    input_filename: str,
    output_filename: str,
    ops: list[dict[str, Any]],
    stop_on_error: bool = False,
) -> dict[str, Any]:
    """Apply a sequence of editing ops to the document in a single open/save.

    Each op is ``{"op": "<name>", **op_kwargs}``. Supported ops:

        - ``replace`` — search_and_replace; kwargs ``find/replace/scope/regex/lenient``
        - ``replace_many`` — kwargs ``replacements: [{find,replace,…}, …]``
        - ``fill_placeholders`` — kwargs ``mapping/pattern``
        - ``insert`` — kwargs ``text/position/anchor/style/format`` (``format`` auto-renders markdown by default)
        - ``insert_image`` — kwargs ``image_filename/position/anchor/width_inches/width_cm/alignment`` (image must already exist in workdir)
        - ``format`` — kwargs ``paragraph_index/paragraph_indexes/anchor/style_filter/bold/...``
        - ``replace_paragraph`` — kwargs ``anchor/new_text/style/format`` (multi-line ``new_text`` auto-splits; ``format`` auto-renders markdown)
        - ``replace_section`` — kwargs ``heading_anchor/new_content/preserve_heading/style/format`` (rewrite a whole section atomically; ``format`` auto-renders markdown)
        - ``delete_paragraph`` — kwargs ``anchor``
        - ``delete_range`` — kwargs ``start_anchor/end_anchor/include_end`` (delete a contiguous block atomically)
        - ``set_cell_text`` — kwargs ``table_index/row/col/text/preserve_format``
        - ``fill_table`` — kwargs ``table_index/rows/mode/has_header``
        - ``add_table`` — kwargs ``rows|markdown/has_header/caption/position/anchor``
        - ``update_field`` — kwargs ``field/value`` (core docprops: TITLE/AUTHOR/SUBJECT/KEYWORDS/DESCRIPTION/CATEGORY)

    Pre-execution safety:
        Consecutive ``delete_paragraph`` ops with INTEGER anchors are auto-
        reordered into descending anchor order before execution. This avoids
        the classic index-drift bug where deleting paragraph N shifts later
        paragraphs down by one, causing every other paragraph in the run to
        be skipped. Non-integer (text) anchors and non-delete ops are left
        in original order.

    Returns a per-op result list. By default a failed op is logged but does
    NOT abort the batch (``stop_on_error=False``); pass True to raise on
    first failure (the document is saved up through the prior successful op).
    """
    if not isinstance(ops, list) or not ops:
        raise ValueError("'ops' must be a non-empty list of {op, …} dicts")

    ops, reordered_count = _normalize_int_anchor_deletes(ops)

    doc = _open(input_filename)
    results: list[dict[str, Any]] = []
    succeeded = 0
    failed = 0

    for i, op_def in enumerate(ops):
        if not isinstance(op_def, dict):
            entry = {"index": i, "ok": False, "error": "op must be a dict"}
            results.append(entry)
            failed += 1
            if stop_on_error:
                break
            continue

        op_name = op_def.get("op")
        op_kwargs = {k: v for k, v in op_def.items() if k != "op"}
        fn = _OP_REGISTRY.get(op_name) if isinstance(op_name, str) else None
        if fn is None:
            entry = {
                "index": i, "op": op_name, "ok": False,
                "error": f"unknown op {op_name!r}; supported: {sorted(_OP_REGISTRY.keys())}",
            }
            results.append(entry)
            failed += 1
            if stop_on_error:
                break
            continue

        try:
            sub = fn(doc, **op_kwargs)
            entry = {"index": i, "op": op_name, "ok": True, **sub}
            succeeded += 1
        except Exception as exc:
            entry = {
                "index": i, "op": op_name, "ok": False,
                "error": f"{type(exc).__name__}: {exc}",
            }
            failed += 1
        results.append(entry)
        if not entry["ok"] and stop_on_error:
            break

    _save(doc, output_filename)
    response: dict[str, Any] = {
        "output_filename": output_filename,
        "ops_total": len(ops),
        "ops_succeeded": succeeded,
        "ops_failed": failed,
        "results": results,
    }
    if reordered_count > 0:
        response["int_anchor_deletes_reordered"] = reordered_count
    return response
